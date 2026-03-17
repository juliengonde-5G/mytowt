"""
Commercial Pricing Router — Rate Grids, Clients, Rate Offers.
Handles tariff management for Freight Forwarders and Shippers.
"""
from fastapi import APIRouter, Request, Depends, Form, HTTPException, Query
from fastapi.responses import HTMLResponse, RedirectResponse, FileResponse
from app.templating import templates
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func, and_
from sqlalchemy.orm import selectinload
from typing import Optional
from datetime import datetime, date
import os, json

from app.database import get_db
from app.auth import get_current_user
from app.models.user import User
from app.models.vessel import Vessel
from app.models.port import Port
from app.models.leg import Leg
from app.models.order import Order
from app.models.finance import OpexParameter
from app.models.commercial import (
    Client, ClientType, RateGrid, RateGridStatus, RateGridLine,
    RateOffer, RateOfferStatus, PALETTE_BRACKETS,
    DEFAULT_BRACKETS_SHIPPER, DEFAULT_BRACKETS_FF
)
from app.utils.navigation import haversine_nm, compute_nav_days
from app.utils.activity import log_activity, get_client_ip

router = APIRouter(prefix="/commercial/pricing", tags=["pricing"])

OFFER_DIR = "app/static/uploads/rate_offers"
OPEX_DAILY_DEFAULT = 11600


def _generate_offer_docx(path, reference, client_name, client_contact,
                          valid_from, valid_to, validity_date,
                          bl_fee, booking_fee, notes, lines):
    """Generate a DOCX rate offer document using python-docx."""
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Header
    h = doc.add_heading("TOWT — Offre tarifaire", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x09, 0x55, 0x61)

    doc.add_paragraph(f"Référence : {reference}")
    doc.add_paragraph(f"Client : {client_name}")
    if client_contact:
        doc.add_paragraph(f"Contact : {client_contact}")
    doc.add_paragraph(f"Période de validité : {valid_from} → {valid_to}")
    doc.add_paragraph(f"Date limite de l'offre : {validity_date}")
    doc.add_paragraph("")

    # Fees
    doc.add_heading("Frais", level=2)
    doc.add_paragraph(f"BL Fee : {bl_fee:.2f} €")
    doc.add_paragraph(f"Booking Fee : {booking_fee:.2f} €")
    doc.add_paragraph("")

    # Rate table
    doc.add_heading("Tarifs par route (€ / palette)", level=2)

    if lines:
        headers = ["POL", "POD", "Dist. (NM)", "Jours nav.",
                    "< 10 pal.", "10-50 pal.", "51-100 pal.", "> 100 pal."]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        hdr = table.rows[0]
        for i, text in enumerate(headers):
            cell = hdr.cells[i]
            cell.text = text
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            from docx.oxml.ns import qn
            shading = cell._tc.get_or_add_tcPr()
            bg = shading.makeelement(qn("w:shd"), {
                qn("w:fill"): "095561", qn("w:val"): "clear",
            })
            shading.append(bg)

        # Data rows
        for line in lines:
            row = table.add_row()
            vals = [
                f"{line['pol_name']} ({line['pol_locode']})",
                f"{line['pod_name']} ({line['pod_locode']})",
                f"{line['distance_nm']:.0f}" if line.get("distance_nm") else "—",
                f"{line['nav_days']:.1f}" if line.get("nav_days") else "—",
                f"{line['rate_lt10']:.2f}" if line.get("rate_lt10") else "—",
                f"{line['rate_10to50']:.2f}" if line.get("rate_10to50") else "—",
                f"{line['rate_51to100']:.2f}" if line.get("rate_51to100") else "—",
                f"{line['rate_gt100']:.2f}" if line.get("rate_gt100") else "—",
            ]
            for i, val in enumerate(vals):
                cell = row.cells[i]
                cell.text = val
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.size = Pt(9)

    if notes:
        doc.add_paragraph("")
        doc.add_heading("Notes", level=2)
        doc.add_paragraph(notes)

    # Footer
    doc.add_paragraph("")
    p = doc.add_paragraph("TOWT — Transport à la Voile")
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    p = doc.add_paragraph("Les Docks, 52 quai Frissard — 76600 Le Havre")
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(path)


# ─── Helpers ────────────────────────────────────────────────
def pf(val, default=None):
    if val is None or (isinstance(val, str) and val.strip() == ""):
        return default
    try:
        if isinstance(val, str):
            val = val.replace(" ", "").replace("\u00a0", "")
            if "." in val and "," in val:
                val = val.replace(".", "").replace(",", ".")
            elif "," in val:
                val = val.replace(",", ".")
        return float(val)
    except:
        return default


def pi(val, default=None):
    if val is None or (isinstance(val, str) and val.strip() == ""):
        return default
    try:
        return int(val)
    except:
        return default


def pd(val):
    if val is None or (isinstance(val, str) and val.strip() == ""):
        return None
    try:
        return date.fromisoformat(val)
    except:
        return None


async def get_opex_daily(db: AsyncSession) -> float:
    result = await db.execute(
        select(OpexParameter).where(OpexParameter.parameter_name == "opex_daily_rate")
    )
    param = result.scalar_one_or_none()
    return param.parameter_value if param else OPEX_DAILY_DEFAULT


async def get_company_param(db: AsyncSession, name: str, default: float = 0) -> float:
    result = await db.execute(
        select(OpexParameter).where(OpexParameter.parameter_name == name)
    )
    param = result.scalar_one_or_none()
    return param.parameter_value if param else default


async def generate_grid_ref(db: AsyncSession) -> str:
    year = datetime.now().year
    prefix = f"RG-{year}-"
    result = await db.execute(
        select(func.count(RateGrid.id)).where(RateGrid.reference.like(f"{prefix}%"))
    )
    count = result.scalar() or 0
    return f"{prefix}{count + 1:04d}"


async def generate_offer_ref(db: AsyncSession) -> str:
    year = datetime.now().year
    prefix = f"RO-{year}-"
    result = await db.execute(
        select(func.count(RateOffer.id)).where(RateOffer.reference.like(f"{prefix}%"))
    )
    count = result.scalar() or 0
    return f"{prefix}{count + 1:04d}"


async def get_port_distance(db: AsyncSession, pol_locode: str, pod_locode: str) -> float:
    """Get orthodromic distance between two ports."""
    pol_r = await db.execute(select(Port).where(Port.locode == pol_locode))
    pod_r = await db.execute(select(Port).where(Port.locode == pod_locode))
    pol = pol_r.scalar_one_or_none()
    pod = pod_r.scalar_one_or_none()
    if not pol or not pod or not pol.latitude or not pod.latitude:
        return 0
    return round(haversine_nm(pol.latitude, pol.longitude, pod.latitude, pod.longitude), 1)


# ═══════════════════════════════════════════════════════════
# CLIENTS
# ═══════════════════════════════════════════════════════════

@router.get("/clients", response_class=HTMLResponse)
async def clients_list(
    request: Request,
    client_type: Optional[str] = Query(None),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    query = select(Client).order_by(Client.name)
    if client_type:
        query = query.where(Client.client_type == client_type)
    result = await db.execute(query)
    clients = result.scalars().all()
    return templates.TemplateResponse("commercial/clients.html", {
        "request": request, "user": user, "clients": clients,
        "selected_type": client_type,
        "active_module": "commercial",
    })


@router.get("/clients/create", response_class=HTMLResponse)
async def client_create_form(
    request: Request,
    user: User = Depends(get_current_user),
):
    from app.utils.pipedrive import is_configured
    return templates.TemplateResponse("commercial/client_form.html", {
        "request": request, "user": user,
        "edit_client": None,
        "pipedrive_enabled": await is_configured(),
        "active_module": "commercial",
    })


@router.post("/clients/create", response_class=HTMLResponse)
async def client_create_submit(
    request: Request,
    name: str = Form(...),
    client_type: str = Form(...),
    contact_name: Optional[str] = Form(None),
    contact_email: Optional[str] = Form(None),
    contact_phone: Optional[str] = Form(None),
    address: Optional[str] = Form(None),
    country: Optional[str] = Form(None),
    notes: Optional[str] = Form(None),
    pipedrive_org_id: Optional[int] = Form(None),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    client = Client(
        name=name.strip(),
        client_type=client_type,
        contact_name=contact_name.strip() if contact_name else None,
        contact_email=contact_email.strip() if contact_email else None,
        contact_phone=contact_phone.strip() if contact_phone else None,
        address=address.strip() if address else None,
        country=country.strip() if country else None,
        notes=notes.strip() if notes else None,
        pipedrive_org_id=pipedrive_org_id,
    )
    db.add(client)
    await db.flush()
    await log_activity(db, user=user, action="create", module="commercial",
                       entity_type="client", entity_id=client.id,
                       entity_label=name.strip(), ip_address=get_client_ip(request))
    url = "/commercial/pricing/clients"
    if request.headers.get("HX-Request"):
        return HTMLResponse(content="", headers={"HX-Redirect": url})
    return RedirectResponse(url=url, status_code=303)


@router.get("/clients/{cid}/edit", response_class=HTMLResponse)
async def client_edit_form(
    cid: int, request: Request,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(select(Client).where(Client.id == cid))
    client = result.scalar_one_or_none()
    if not client:
        raise HTTPException(404)
    from app.utils.pipedrive import is_configured
    return templates.TemplateResponse("commercial/client_form.html", {
        "request": request, "user": user,
        "edit_client": client,
        "pipedrive_enabled": await is_configured(),
        "active_module": "commercial",
    })


@router.post("/clients/{cid}/edit", response_class=HTMLResponse)
async def client_edit_submit(
    cid: int, request: Request,
    name: str = Form(...),
    client_type: str = Form(...),
    contact_name: Optional[str] = Form(None),
    contact_email: Optional[str] = Form(None),
    contact_phone: Optional[str] = Form(None),
    address: Optional[str] = Form(None),
    country: Optional[str] = Form(None),
    notes: Optional[str] = Form(None),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(select(Client).where(Client.id == cid))
    client = result.scalar_one_or_none()
    if not client:
        raise HTTPException(404)
    client.name = name.strip()
    client.client_type = client_type
    client.contact_name = contact_name.strip() if contact_name else None
    client.contact_email = contact_email.strip() if contact_email else None
    client.contact_phone = contact_phone.strip() if contact_phone else None
    client.address = address.strip() if address else None
    client.country = country.strip() if country else None
    client.notes = notes.strip() if notes else None
    url = "/commercial/pricing/clients"
    if request.headers.get("HX-Request"):
        return HTMLResponse(content="", headers={"HX-Redirect": url})
    return RedirectResponse(url=url, status_code=303)


# ═══════════════════════════════════════════════════════════
# PIPEDRIVE INTEGRATION
# ═══════════════════════════════════════════════════════════

@router.get("/pipedrive/search")
async def pipedrive_search(
    request: Request,
    q: str = Query("", min_length=2),
    user: User = Depends(get_current_user),
):
    """Search Pipedrive Organizations (HTMX endpoint, returns JSON)."""
    from app.utils.pipedrive import search_organizations
    results = await search_organizations(q, limit=8)
    from fastapi.responses import JSONResponse
    return JSONResponse(content=results)


@router.get("/pipedrive/org/{org_id}")
async def pipedrive_get_org(
    org_id: int, request: Request,
    user: User = Depends(get_current_user),
):
    """Get full Organization details from Pipedrive (returns JSON)."""
    from app.utils.pipedrive import get_organization
    org = await get_organization(org_id)
    if not org:
        raise HTTPException(404, "Organisation non trouvée dans Pipedrive")
    from fastapi.responses import JSONResponse
    return JSONResponse(content=org)


# ═══════════════════════════════════════════════════════════
# INLINE CLIENT CREATION (HTMX from grid form)
# ═══════════════════════════════════════════════════════════

@router.post("/clients/create-inline", response_class=HTMLResponse)
async def client_create_inline(
    request: Request,
    name: str = Form(...),
    client_type: str = Form(...),
    contact_name: Optional[str] = Form(None),
    contact_email: Optional[str] = Form(None),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Create a client inline (from grid form) and return updated select options."""
    client = Client(
        name=name.strip(),
        client_type=client_type,
        contact_name=contact_name.strip() if contact_name else None,
        contact_email=contact_email.strip() if contact_email else None,
    )
    db.add(client)
    await db.flush()

    # Return updated client select options
    clients_r = await db.execute(select(Client).where(Client.is_active == True).order_by(Client.name))
    clients = clients_r.scalars().all()

    html = '<option value="">— Sélectionner —</option>'
    for c in clients:
        selected = ' selected' if c.id == client.id else ''
        label = "FF" if c.client_type == "freight_forwarder" else "Chargeur"
        html += f'<option value="{c.id}" data-type="{c.client_type}"{selected}>{c.name} ({label})</option>'

    return HTMLResponse(content=html)


# ═══════════════════════════════════════════════════════════
# API: Get sailing schedule legs for a date range
# ═══════════════════════════════════════════════════════════

@router.get("/api/legs-for-period", response_class=HTMLResponse)
async def api_legs_for_period(
    request: Request,
    valid_from: str = Query(...),
    valid_to: str = Query(...),
    vessel_id: Optional[int] = Query(None),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Return unique POL/POD routes from legs within the given period."""
    from_date = pd(valid_from)
    to_date = pd(valid_to)
    if not from_date or not to_date:
        return HTMLResponse(content="")

    # Get year(s) from the period
    years = list(range(from_date.year, to_date.year + 1))

    query = select(Leg).options(
        selectinload(Leg.departure_port),
        selectinload(Leg.arrival_port),
    ).where(Leg.year.in_(years)).order_by(Leg.vessel_id, Leg.sequence)

    if vessel_id:
        query = query.where(Leg.vessel_id == vessel_id)

    result = await db.execute(query)
    legs = result.scalars().all()

    # Build unique POL/POD routes only
    seen_routes = set()
    routes = []
    for leg in legs:
        route_key = f"{leg.departure_port_locode}-{leg.arrival_port_locode}"
        if route_key not in seen_routes:
            seen_routes.add(route_key)
            dep_name = leg.departure_port.name if leg.departure_port else leg.departure_port_locode
            arr_name = leg.arrival_port.name if leg.arrival_port else leg.arrival_port_locode
            routes.append({
                "pol": leg.departure_port_locode,
                "pod": leg.arrival_port_locode,
                "pol_name": dep_name,
                "pod_name": arr_name,
            })

    import json as json_mod
    return HTMLResponse(
        content=json_mod.dumps(routes),
        media_type="application/json"
    )


# ═══════════════════════════════════════════════════════════
# API: Compute rates for a set of routes
# ═══════════════════════════════════════════════════════════

@router.post("/api/compute-rates", response_class=HTMLResponse)
async def api_compute_rates(
    request: Request,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Compute rates for given routes, OPEX, brackets, adjustment index."""
    import json as json_mod
    body = await request.json()

    vessel_id = body.get("vessel_id")
    adjustment_index = float(body.get("adjustment_index", 1.0))
    brackets = body.get("brackets", DEFAULT_BRACKETS_SHIPPER)
    routes = body.get("routes", [])

    # Get OPEX
    opex_daily = OPEX_DAILY_DEFAULT
    if vessel_id:
        v_r = await db.execute(select(Vessel).where(Vessel.id == int(vessel_id)))
        vessel = v_r.scalar_one_or_none()
        if vessel and vessel.opex_daily_sea:
            opex_daily = vessel.opex_daily_sea
    if not opex_daily or opex_daily == 0:
        opex_daily = await get_opex_daily(db)

    results = []
    for route in routes:
        pol = route.get("pol", "").upper()
        pod = route.get("pod", "").upper()
        distance = await get_port_distance(db, pol, pod)
        nav_days = round(distance / (8 * 24), 2) if distance else 0
        base_rate = round(opex_daily * nav_days / 850, 2) if nav_days else 0

        rates = {}
        for b in brackets:
            rates[b["key"]] = round(base_rate * b["coeff"] * adjustment_index, 2)

        results.append({
            "pol": pol,
            "pod": pod,
            "distance": round(distance, 0) if distance else 0,
            "nav_days": nav_days,
            "base_rate": base_rate,
            "opex_daily": opex_daily,
            "rates": rates,
        })

    return HTMLResponse(
        content=json_mod.dumps(results),
        media_type="application/json"
    )


# ═══════════════════════════════════════════════════════════
# API: Get default brackets for client type
# ═══════════════════════════════════════════════════════════

@router.get("/api/default-brackets", response_class=HTMLResponse)
async def api_default_brackets(
    request: Request,
    client_type: str = Query("shipper"),
    user: User = Depends(get_current_user),
):
    import json as json_mod
    if client_type == "freight_forwarder":
        return HTMLResponse(content=json_mod.dumps(DEFAULT_BRACKETS_FF), media_type="application/json")
    return HTMLResponse(content=json_mod.dumps(DEFAULT_BRACKETS_SHIPPER), media_type="application/json")


# ═══════════════════════════════════════════════════════════
# RATE GRIDS
# ═══════════════════════════════════════════════════════════

@router.get("/grids", response_class=HTMLResponse)
async def grids_list(
    request: Request,
    client_id: Optional[int] = Query(None),
    status: Optional[str] = Query(None),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    query = select(RateGrid).options(
        selectinload(RateGrid.client),
        selectinload(RateGrid.vessel),
        selectinload(RateGrid.lines).selectinload(RateGridLine.pol),
        selectinload(RateGrid.lines).selectinload(RateGridLine.pod),
    ).order_by(RateGrid.created_at.desc())
    if client_id:
        query = query.where(RateGrid.client_id == client_id)
    if status:
        query = query.where(RateGrid.status == status)
    result = await db.execute(query)
    grids = result.scalars().all()

    clients_r = await db.execute(select(Client).where(Client.is_active == True).order_by(Client.name))
    clients = clients_r.scalars().all()

    return templates.TemplateResponse("commercial/grids.html", {
        "request": request, "user": user,
        "grids": grids, "clients": clients,
        "selected_client_id": client_id, "selected_status": status,
        "active_module": "commercial",
    })


@router.get("/grids/create", response_class=HTMLResponse)
async def grid_create_form(
    request: Request,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    ref = await generate_grid_ref(db)
    clients_r = await db.execute(select(Client).where(Client.is_active == True).order_by(Client.name))
    clients = clients_r.scalars().all()
    vessels_r = await db.execute(select(Vessel).where(Vessel.is_active == True).order_by(Vessel.name))
    vessels = vessels_r.scalars().all()
    ports_r = await db.execute(select(Port).where(Port.is_shortcut == True).order_by(Port.name))
    ports = ports_r.scalars().all()

    bl_fee = await get_company_param(db, "bl_fee", 0)
    booking_fee = await get_company_param(db, "booking_fee", 0)

    import json as json_mod
    return templates.TemplateResponse("commercial/grid_form.html", {
        "request": request, "user": user,
        "edit_grid": None, "reference": ref,
        "clients": clients, "vessels": vessels, "ports": ports,
        "default_bl_fee": bl_fee, "default_booking_fee": booking_fee,
        "default_brackets_shipper": json_mod.dumps(DEFAULT_BRACKETS_SHIPPER),
        "default_brackets_ff": json_mod.dumps(DEFAULT_BRACKETS_FF),
        "active_module": "commercial",
    })


@router.post("/grids/create", response_class=HTMLResponse)
async def grid_create_submit(
    request: Request,
    client_id: str = Form(...),
    vessel_id: Optional[str] = Form(None),
    valid_from: str = Form(...),
    valid_to: str = Form(...),
    adjustment_index: Optional[str] = Form(None),
    bl_fee: Optional[str] = Form(None),
    booking_fee: Optional[str] = Form(None),
    volume_commitment: Optional[str] = Form(None),
    brackets_json: Optional[str] = Form(None),
    notes: Optional[str] = Form(None),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    import json as json_mod

    ref = await generate_grid_ref(db)

    # Load client to determine type
    cid = pi(client_id)
    client_r = await db.execute(select(Client).where(Client.id == cid))
    client = client_r.scalar_one_or_none()

    grid = RateGrid(
        reference=ref,
        client_id=cid,
        vessel_id=pi(vessel_id),
        valid_from=pd(valid_from),
        valid_to=pd(valid_to),
        adjustment_index=pf(adjustment_index, 1.0),
        bl_fee=pf(bl_fee, 0),
        booking_fee=pf(booking_fee, 0),
        volume_commitment=pi(volume_commitment),
        brackets_json=brackets_json if brackets_json else None,
        notes=notes.strip() if notes else None,
        created_by=user.username,
    )
    db.add(grid)
    await db.flush()

    # Parse brackets
    brackets = grid.brackets

    # Get OPEX daily rate from vessel or company default
    opex_daily = OPEX_DAILY_DEFAULT
    if grid.vessel_id:
        v_r = await db.execute(select(Vessel).where(Vessel.id == grid.vessel_id))
        vessel = v_r.scalar_one_or_none()
        if vessel and vessel.opex_daily_sea:
            opex_daily = vessel.opex_daily_sea
    if not opex_daily or opex_daily == 0:
        opex_daily = await get_opex_daily(db)

    # Parse route lines from form
    form_data = await request.form()
    i = 0
    while True:
        pol_key = f"pol_{i}"
        pod_key = f"pod_{i}"
        if pol_key not in form_data or pod_key not in form_data:
            break
        pol = form_data[pol_key].strip().upper() if form_data[pol_key] else None
        pod = form_data[pod_key].strip().upper() if form_data[pod_key] else None
        leg_id_val = pi(form_data.get(f"leg_id_{i}"))

        if pol and pod:
            distance = await get_port_distance(db, pol, pod)

            # Check for manual rate overrides
            manual_rates = {}
            has_manual = False
            for b in brackets:
                manual_val = pf(form_data.get(f"rate_{b['key']}_{i}"))
                if manual_val is not None:
                    manual_rates[b["key"]] = manual_val
                    has_manual = True

            line = RateGridLine(
                rate_grid_id=grid.id,
                pol_locode=pol,
                pod_locode=pod,
                leg_id=leg_id_val,
                distance_nm=distance,
                is_manual=has_manual,
            )

            if has_manual:
                line.opex_daily = opex_daily
                line.nav_days = compute_nav_days(distance)
                line.base_rate = round(opex_daily * (line.nav_days or 0) / 850, 2) if line.nav_days else None
                line.rates = manual_rates
            else:
                line.compute_rates(opex_daily, grid.adjustment_index, brackets)

            db.add(line)
        i += 1

    await log_activity(db, user=user, action="create", module="commercial",
                       entity_type="rate_grid", entity_id=grid.id,
                       entity_label=ref, ip_address=get_client_ip(request))

    url = f"/commercial/pricing/grids/{grid.id}"
    if request.headers.get("HX-Request"):
        return HTMLResponse(content="", headers={"HX-Redirect": url})
    return RedirectResponse(url=url, status_code=303)


@router.get("/grids/{gid}", response_class=HTMLResponse)
async def grid_detail(
    gid: int, request: Request,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(RateGrid).options(
            selectinload(RateGrid.client),
            selectinload(RateGrid.vessel),
            selectinload(RateGrid.lines).selectinload(RateGridLine.pol),
            selectinload(RateGrid.lines).selectinload(RateGridLine.pod),
        ).where(RateGrid.id == gid)
    )
    grid = result.scalar_one_or_none()
    if not grid:
        raise HTTPException(404)

    # Get related offers
    offers_r = await db.execute(
        select(RateOffer).where(RateOffer.rate_grid_id == gid).order_by(RateOffer.created_at.desc())
    )
    offers = offers_r.scalars().all()

    # Get orders linked to this grid
    orders_r = await db.execute(
        select(Order).where(Order.rate_grid_id == gid).order_by(Order.created_at.desc())
    )
    orders = orders_r.scalars().all()

    return templates.TemplateResponse("commercial/grid_detail.html", {
        "request": request, "user": user,
        "grid": grid, "offers": offers, "orders": orders,
        "brackets": PALETTE_BRACKETS,
        "active_module": "commercial",
    })


@router.get("/grids/{gid}/edit", response_class=HTMLResponse)
async def grid_edit_form(
    gid: int, request: Request,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(RateGrid).options(
            selectinload(RateGrid.client),
            selectinload(RateGrid.vessel),
            selectinload(RateGrid.lines).selectinload(RateGridLine.pol),
            selectinload(RateGrid.lines).selectinload(RateGridLine.pod),
        ).where(RateGrid.id == gid)
    )
    grid = result.scalar_one_or_none()
    if not grid:
        raise HTTPException(404)

    clients_r = await db.execute(select(Client).where(Client.is_active == True).order_by(Client.name))
    clients = clients_r.scalars().all()
    vessels_r = await db.execute(select(Vessel).where(Vessel.is_active == True).order_by(Vessel.name))
    vessels = vessels_r.scalars().all()
    ports_r = await db.execute(select(Port).where(Port.is_shortcut == True).order_by(Port.name))
    ports = ports_r.scalars().all()

    return templates.TemplateResponse("commercial/grid_form.html", {
        "request": request, "user": user,
        "edit_grid": grid, "reference": grid.reference,
        "clients": clients, "vessels": vessels, "ports": ports,
        "default_bl_fee": grid.bl_fee, "default_booking_fee": grid.booking_fee,
        "active_module": "commercial",
    })


@router.post("/grids/{gid}/edit", response_class=HTMLResponse)
async def grid_edit_submit(
    gid: int, request: Request,
    client_id: str = Form(...),
    vessel_id: Optional[str] = Form(None),
    valid_from: str = Form(...),
    valid_to: str = Form(...),
    adjustment_index: Optional[str] = Form(None),
    bl_fee: Optional[str] = Form(None),
    booking_fee: Optional[str] = Form(None),
    volume_commitment: Optional[str] = Form(None),
    status: Optional[str] = Form(None),
    notes: Optional[str] = Form(None),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(RateGrid).options(selectinload(RateGrid.lines)).where(RateGrid.id == gid)
    )
    grid = result.scalar_one_or_none()
    if not grid:
        raise HTTPException(404)

    grid.client_id = pi(client_id, grid.client_id)
    grid.vessel_id = pi(vessel_id)
    grid.valid_from = pd(valid_from) or grid.valid_from
    grid.valid_to = pd(valid_to) or grid.valid_to
    grid.adjustment_index = pf(adjustment_index, 1.0)
    grid.bl_fee = pf(bl_fee, 0)
    grid.booking_fee = pf(booking_fee, 0)
    grid.volume_commitment = pi(volume_commitment)
    grid.notes = notes.strip() if notes else None
    if status:
        grid.status = status

    # Get OPEX
    opex_daily = OPEX_DAILY_DEFAULT
    if grid.vessel_id:
        v_r = await db.execute(select(Vessel).where(Vessel.id == grid.vessel_id))
        vessel = v_r.scalar_one_or_none()
        if vessel and vessel.opex_daily_sea:
            opex_daily = vessel.opex_daily_sea
    if not opex_daily or opex_daily == 0:
        opex_daily = await get_opex_daily(db)

    # Remove old lines
    for old_line in list(grid.lines):
        await db.delete(old_line)
    await db.flush()

    # Re-add lines from form
    form_data = await request.form()
    i = 0
    while True:
        pol_key = f"pol_{i}"
        pod_key = f"pod_{i}"
        if pol_key not in form_data or pod_key not in form_data:
            break
        pol = form_data[pol_key].strip().upper() if form_data[pol_key] else None
        pod = form_data[pod_key].strip().upper() if form_data[pod_key] else None

        # Check for manual overrides
        manual_lt10 = pf(form_data.get(f"rate_lt10_{i}"))
        manual_10to50 = pf(form_data.get(f"rate_10to50_{i}"))
        manual_51to100 = pf(form_data.get(f"rate_51to100_{i}"))
        manual_gt100 = pf(form_data.get(f"rate_gt100_{i}"))
        is_manual = any(v is not None for v in [manual_lt10, manual_10to50, manual_51to100, manual_gt100])

        if pol and pod:
            distance = await get_port_distance(db, pol, pod)
            line = RateGridLine(
                rate_grid_id=grid.id,
                pol_locode=pol,
                pod_locode=pod,
                distance_nm=distance,
                is_manual=is_manual,
            )
            if is_manual:
                line.opex_daily = opex_daily
                line.nav_days = compute_nav_days(distance)
                line.base_rate = round(opex_daily * (line.nav_days or 0) / 850, 2) if line.nav_days else None
                line.rate_lt10 = manual_lt10
                line.rate_10to50 = manual_10to50
                line.rate_51to100 = manual_51to100
                line.rate_gt100 = manual_gt100
            else:
                line.compute_rates(opex_daily, grid.adjustment_index)
            db.add(line)
        i += 1

    url = f"/commercial/pricing/grids/{gid}"
    if request.headers.get("HX-Request"):
        return HTMLResponse(content="", headers={"HX-Redirect": url})
    return RedirectResponse(url=url, status_code=303)


@router.post("/grids/{gid}/delete", response_class=HTMLResponse)
async def grid_delete(
    gid: int, request: Request,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(RateGrid).options(selectinload(RateGrid.lines)).where(RateGrid.id == gid)
    )
    grid = result.scalar_one_or_none()
    if not grid:
        raise HTTPException(404)

    # Cannot delete active grids
    if grid.status == RateGridStatus.ACTIVE.value:
        return HTMLResponse("Impossible de supprimer une grille active.", status_code=400)

    ref = grid.reference
    # Delete lines first
    for line in grid.lines:
        await db.delete(line)
    await db.delete(grid)

    await log_activity(db, user=user, action="delete", module="commercial",
                       entity_type="rate_grid", entity_id=gid,
                       entity_label=ref, ip_address=get_client_ip(request))
    await db.commit()

    if request.headers.get("HX-Request"):
        return HTMLResponse(content="", headers={"HX-Redirect": "/commercial/pricing/grids"})
    return RedirectResponse(url="/commercial/pricing/grids", status_code=303)


@router.post("/grids/{gid}/activate", response_class=HTMLResponse)
async def grid_activate(
    gid: int, request: Request,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(select(RateGrid).where(RateGrid.id == gid))
    grid = result.scalar_one_or_none()
    if not grid:
        raise HTTPException(404)

    # Supersede other active grids for same client
    active_r = await db.execute(
        select(RateGrid).where(
            and_(RateGrid.client_id == grid.client_id,
                 RateGrid.status == RateGridStatus.ACTIVE.value,
                 RateGrid.id != gid)
        )
    )
    for old in active_r.scalars().all():
        old.status = RateGridStatus.SUPERSEDED.value

    grid.status = RateGridStatus.ACTIVE.value
    await log_activity(db, user=user, action="update", module="commercial",
                       entity_type="rate_grid", entity_id=gid,
                       entity_label=grid.reference, detail="Activated",
                       ip_address=get_client_ip(request))
    url = f"/commercial/pricing/grids/{gid}"
    if request.headers.get("HX-Request"):
        return HTMLResponse(content="", headers={"HX-Redirect": url})
    return RedirectResponse(url=url, status_code=303)


@router.post("/grids/{gid}/recalculate", response_class=HTMLResponse)
async def grid_recalculate(
    gid: int, request: Request,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Recalculate all non-manual lines in a grid."""
    result = await db.execute(
        select(RateGrid).options(selectinload(RateGrid.lines)).where(RateGrid.id == gid)
    )
    grid = result.scalar_one_or_none()
    if not grid:
        raise HTTPException(404)

    opex_daily = OPEX_DAILY_DEFAULT
    if grid.vessel_id:
        v_r = await db.execute(select(Vessel).where(Vessel.id == grid.vessel_id))
        vessel = v_r.scalar_one_or_none()
        if vessel and vessel.opex_daily_sea:
            opex_daily = vessel.opex_daily_sea
    if not opex_daily or opex_daily == 0:
        opex_daily = await get_opex_daily(db)

    for line in grid.lines:
        if not line.is_manual:
            if not line.distance_nm:
                line.distance_nm = await get_port_distance(db, line.pol_locode, line.pod_locode)
            line.compute_rates(opex_daily, grid.adjustment_index)

    url = f"/commercial/pricing/grids/{gid}"
    if request.headers.get("HX-Request"):
        return HTMLResponse(content="", headers={"HX-Redirect": url})
    return RedirectResponse(url=url, status_code=303)


# ═══════════════════════════════════════════════════════════
# RATE OFFERS (DOCX generation)
# ═══════════════════════════════════════════════════════════

@router.get("/grids/{gid}/offer/create", response_class=HTMLResponse)
async def offer_create_form(
    gid: int, request: Request,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(RateGrid).options(
            selectinload(RateGrid.client),
            selectinload(RateGrid.lines).selectinload(RateGridLine.pol),
            selectinload(RateGrid.lines).selectinload(RateGridLine.pod),
        ).where(RateGrid.id == gid)
    )
    grid = result.scalar_one_or_none()
    if not grid:
        raise HTTPException(404)

    ref = await generate_offer_ref(db)

    return templates.TemplateResponse("commercial/offer_form.html", {
        "request": request, "user": user,
        "grid": grid, "reference": ref,
        "active_module": "commercial",
    })


@router.post("/grids/{gid}/offer/create", response_class=HTMLResponse)
async def offer_create_submit(
    gid: int, request: Request,
    validity_date: str = Form(...),
    notes: Optional[str] = Form(None),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(RateGrid).options(
            selectinload(RateGrid.client),
            selectinload(RateGrid.lines).selectinload(RateGridLine.pol),
            selectinload(RateGrid.lines).selectinload(RateGridLine.pod),
        ).where(RateGrid.id == gid)
    )
    grid = result.scalar_one_or_none()
    if not grid:
        raise HTTPException(404)

    ref = await generate_offer_ref(db)
    offer = RateOffer(
        reference=ref,
        client_id=grid.client_id,
        rate_grid_id=grid.id,
        validity_date=pd(validity_date),
        notes=notes.strip() if notes else None,
        created_by=user.username,
    )
    db.add(offer)
    await db.flush()

    # Generate DOCX using python-docx
    os.makedirs(OFFER_DIR, exist_ok=True)
    docx_filename = f"rate_offer_{ref.replace('-', '_')}.docx"
    docx_path = os.path.join(OFFER_DIR, docx_filename)

    try:
        _generate_offer_docx(
            path=docx_path,
            reference=ref,
            client_name=grid.client.name if grid.client else "",
            client_contact=grid.client.contact_name if grid.client else "",
            valid_from=str(grid.valid_from) if grid.valid_from else "",
            valid_to=str(grid.valid_to) if grid.valid_to else "",
            validity_date=str(offer.validity_date) if offer.validity_date else "",
            bl_fee=grid.bl_fee or 0,
            booking_fee=grid.booking_fee or 0,
            notes=offer.notes or "",
            lines=[{
                "pol_name": l.pol.name if l.pol else l.pol_locode,
                "pod_name": l.pod.name if l.pod else l.pod_locode,
                "pol_locode": l.pol_locode,
                "pod_locode": l.pod_locode,
                "distance_nm": l.distance_nm,
                "nav_days": l.nav_days,
                "rate_lt10": l.rate_lt10,
                "rate_10to50": l.rate_10to50,
                "rate_51to100": l.rate_51to100,
                "rate_gt100": l.rate_gt100,
            } for l in grid.lines],
        )
        if os.path.exists(docx_path):
            offer.document_filename = docx_filename
            offer.document_path = docx_path
    except Exception as e:
        print(f"DOCX generation exception: {e}")

    await log_activity(db, user=user, action="create", module="commercial",
                       entity_type="rate_offer", entity_id=offer.id,
                       entity_label=ref, ip_address=get_client_ip(request))

    # ─── Push to Pipedrive as Deal ───
    if grid.client and grid.client.pipedrive_org_id:
        try:
            from app.utils.pipedrive import create_deal
            routes = ", ".join(f"{l.pol_locode}→{l.pod_locode}" for l in grid.lines)
            deal_id = await create_deal(
                title=f"Offre {ref} — {grid.client.name}",
                org_id=grid.client.pipedrive_org_id,
                notes=(
                    f"<b>Offre tarifaire {ref}</b><br>"
                    f"Client : {grid.client.name}<br>"
                    f"Grille : {grid.reference}<br>"
                    f"Validité : {grid.valid_from} → {grid.valid_to}<br>"
                    f"Routes : {routes}<br>"
                    f"BL fee : {grid.bl_fee}€ · Booking fee : {grid.booking_fee}€<br>"
                    f"Créé par {user.full_name}"
                ),
            )
            if deal_id:
                offer.pipedrive_deal_id = deal_id
                await db.flush()
        except Exception as e:
            print(f"Pipedrive push error (offer): {e}")

    url = f"/commercial/pricing/grids/{gid}"
    if request.headers.get("HX-Request"):
        return HTMLResponse(content="", headers={"HX-Redirect": url})
    return RedirectResponse(url=url, status_code=303)


@router.get("/offers/{oid}/download")
async def offer_download(
    oid: int,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(select(RateOffer).where(RateOffer.id == oid))
    offer = result.scalar_one_or_none()
    if not offer or not offer.document_path or not os.path.exists(offer.document_path):
        raise HTTPException(404)
    return FileResponse(
        offer.document_path,
        filename=offer.document_filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@router.post("/offers/{oid}/status", response_class=HTMLResponse)
async def offer_update_status(
    oid: int, request: Request,
    status: str = Form(...),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(select(RateOffer).where(RateOffer.id == oid))
    offer = result.scalar_one_or_none()
    if not offer:
        raise HTTPException(404)
    offer.status = status
    if status == RateOfferStatus.SENT.value:
        offer.sent_at = datetime.utcnow()

    url = f"/commercial/pricing/grids/{offer.rate_grid_id}"
    if request.headers.get("HX-Request"):
        return HTMLResponse(content="", headers={"HX-Redirect": url})
    return RedirectResponse(url=url, status_code=303)


# ─── GENERATE ORDER FROM OFFER ────────────────────────────
@router.get("/offers/{oid}/create-order", response_class=HTMLResponse)
async def offer_create_order_form(
    oid: int, request: Request,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Show form to create a transport order from an accepted offer."""
    result = await db.execute(
        select(RateOffer).options(
            selectinload(RateOffer.client),
            selectinload(RateOffer.rate_grid).selectinload(RateGrid.lines).selectinload(RateGridLine.pol),
            selectinload(RateOffer.rate_grid).selectinload(RateGrid.lines).selectinload(RateGridLine.pod),
        ).where(RateOffer.id == oid)
    )
    offer = result.scalar_one_or_none()
    if not offer:
        raise HTTPException(404)

    # Generate next OT reference
    year = datetime.now().year
    prefix = f"OT-{year}-"
    count_r = await db.execute(
        select(func.count(Order.id)).where(Order.reference.like(f"{prefix}%"))
    )
    next_ref = f"{prefix}{(count_r.scalar() or 0) + 1:04d}"

    return templates.TemplateResponse("commercial/order_from_offer.html", {
        "request": request, "user": user,
        "offer": offer, "grid": offer.rate_grid,
        "reference": next_ref,
        "active_module": "commercial",
    })


@router.post("/offers/{oid}/create-order", response_class=HTMLResponse)
async def offer_create_order_submit(
    oid: int, request: Request,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Create a transport order from an offer."""
    result = await db.execute(
        select(RateOffer).options(
            selectinload(RateOffer.client),
            selectinload(RateOffer.rate_grid).selectinload(RateGrid.lines).selectinload(RateGridLine.pol),
            selectinload(RateOffer.rate_grid).selectinload(RateGrid.lines).selectinload(RateGridLine.pod),
        ).where(RateOffer.id == oid)
    )
    offer = result.scalar_one_or_none()
    if not offer:
        raise HTTPException(404)

    form = await request.form()
    line_id = pi(form.get("line_id"))
    quantity = pi(form.get("quantity_palettes"), 1)
    palette_format = form.get("palette_format", "EPAL")

    # Find the selected rate grid line
    grid = offer.rate_grid
    selected_line = None
    if line_id:
        selected_line = next((l for l in grid.lines if l.id == line_id), None)

    # Determine unit price based on quantity bracket
    unit_price = 0
    if selected_line:
        if quantity > 100:
            unit_price = selected_line.rate_gt100 or 0
        elif quantity > 50:
            unit_price = selected_line.rate_51to100 or 0
        elif quantity >= 10:
            unit_price = selected_line.rate_10to50 or 0
        else:
            unit_price = selected_line.rate_lt10 or 0

    # Override with manual price if provided
    manual_price = pf(form.get("unit_price"))
    if manual_price is not None:
        unit_price = manual_price

    # Generate reference
    year = datetime.now().year
    prefix = f"OT-{year}-"
    count_r = await db.execute(
        select(func.count(Order.id)).where(Order.reference.like(f"{prefix}%"))
    )
    ref = f"{prefix}{(count_r.scalar() or 0) + 1:04d}"

    order = Order(
        reference=ref,
        client_name=offer.client.name if offer.client else "",
        client_contact=offer.client.contact_name if offer.client else None,
        quantity_palettes=quantity,
        palette_format=palette_format,
        weight_per_palette=pf(form.get("weight_per_palette"), 0.8),
        unit_price=unit_price,
        booking_fee=grid.booking_fee or 0,
        documentation_fee=grid.bl_fee or 0,
        departure_locode=selected_line.pol_locode if selected_line else None,
        arrival_locode=selected_line.pod_locode if selected_line else None,
        description=form.get("description", "").strip() or None,
        rate_grid_id=grid.id,
        rate_grid_line_id=selected_line.id if selected_line else None,
    )
    order.compute_total()
    db.add(order)
    await db.flush()

    await log_activity(db, user=user, action="create", module="commercial",
                       entity_type="order", entity_id=order.id,
                       entity_label=f"Ordre {ref} depuis offre {offer.reference}",
                       ip_address=get_client_ip(request))

    url = "/commercial?tab=orders"
    if request.headers.get("HX-Request"):
        return HTMLResponse(content="", headers={"HX-Redirect": url})
    return RedirectResponse(url=url, status_code=303)


# ═══════════════════════════════════════════════════════════
# RATE OFFERS LIST (History)
# ═══════════════════════════════════════════════════════════

@router.get("/offers", response_class=HTMLResponse)
async def offers_list(
    request: Request,
    client_id: Optional[int] = Query(None),
    status: Optional[str] = Query(None),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    query = select(RateOffer).options(
        selectinload(RateOffer.client),
        selectinload(RateOffer.rate_grid),
    ).order_by(RateOffer.created_at.desc())
    if client_id:
        query = query.where(RateOffer.client_id == client_id)
    if status:
        query = query.where(RateOffer.status == status)
    result = await db.execute(query)
    offers = result.scalars().all()

    clients_r = await db.execute(select(Client).where(Client.is_active == True).order_by(Client.name))
    clients = clients_r.scalars().all()

    return templates.TemplateResponse("commercial/offers.html", {
        "request": request, "user": user,
        "offers": offers, "clients": clients,
        "selected_client_id": client_id, "selected_status": status,
        "active_module": "commercial",
    })


# ═══════════════════════════════════════════════════════════
# API: Get rate for order assignment
# ═══════════════════════════════════════════════════════════

@router.get("/api/rate-lookup", response_class=HTMLResponse)
async def api_rate_lookup(
    request: Request,
    client_id: int = Query(...),
    pol: str = Query(...),
    pod: str = Query(...),
    quantity: int = Query(...),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Find the active rate grid line for a client/route and return pricing."""
    today = date.today()
    result = await db.execute(
        select(RateGrid).options(
            selectinload(RateGrid.lines)
        ).where(
            and_(
                RateGrid.client_id == client_id,
                RateGrid.status == RateGridStatus.ACTIVE.value,
                RateGrid.valid_from <= today,
                RateGrid.valid_to >= today,
            )
        ).order_by(RateGrid.created_at.desc()).limit(1)
    )
    grid = result.scalar_one_or_none()
    if not grid:
        return HTMLResponse(content='<span class="text-orange-500">Aucune grille active pour ce client</span>')

    matching_line = None
    for line in grid.lines:
        if line.pol_locode == pol.upper() and line.pod_locode == pod.upper():
            matching_line = line
            break

    if not matching_line:
        return HTMLResponse(content=f'<span class="text-orange-500">Route {pol}→{pod} non trouvée dans {grid.reference}</span>')

    rate = matching_line.get_rate_for_quantity(quantity)
    total = round(rate * quantity + (grid.bl_fee or 0) + (grid.booking_fee or 0), 2)

    html = f"""
    <div class="bg-blue-50 border border-blue-200 rounded-lg p-3 text-sm">
        <div class="font-semibold text-blue-800">Grille: {grid.reference}</div>
        <div class="mt-1">Prix/palette: <strong>€{rate:.2f}</strong></div>
        <div>Fret: €{rate * quantity:.2f} + BL: €{grid.bl_fee:.2f} + Booking: €{grid.booking_fee:.2f}</div>
        <div class="font-bold text-blue-900 mt-1">Total: €{total:.2f}</div>
        <input type="hidden" name="rate_grid_id" value="{grid.id}">
        <input type="hidden" name="rate_grid_line_id" value="{matching_line.id}">
        <input type="hidden" name="suggested_unit_price" value="{rate}">
        <input type="hidden" name="suggested_bl_fee" value="{grid.bl_fee}">
        <input type="hidden" name="suggested_booking_fee" value="{grid.booking_fee}">
    </div>
    """
    return HTMLResponse(content=html)
