"""On Board module router — ex-Captain, renamed to On Board."""
import io
import json
from datetime import datetime, timezone, date
from typing import Optional
from fastapi import APIRouter, Request, Depends, Query, Form, HTTPException, File, UploadFile
from fastapi.responses import HTMLResponse, RedirectResponse, StreamingResponse, Response
from sqlalchemy.ext.asyncio import AsyncSession
from app.utils.activity import log_activity, get_client_ip
from sqlalchemy import select, func, and_
from sqlalchemy.orm import selectinload

from app.templating import templates
from app.database import get_db
from app.auth import get_current_user
from app.models.user import User
from app.models.vessel import Vessel
from app.models.leg import Leg
from app.models.order import Order
from app.models.crew import CrewMember, CrewAssignment
from app.models.packing_list import PackingList, PackingListBatch
from app.models.onboard import (
    SofEvent, OnboardNotification, CargoDocument, ETAShift, OnboardAttachment,
    CargoDocumentAttachment,
    SOF_EVENT_TYPES, CARGO_DOC_TYPES, ETA_SHIFT_REASONS, ATTACHMENT_CATEGORIES,
)
from app.utils.timezones import get_port_timezone, utc_offset_label, TIMEZONE_CHOICES

router = APIRouter(prefix="/onboard", tags=["onboard"])


# ═══════════════════════════════════════════════════════════════
#  MAIN PAGE
# ═══════════════════════════════════════════════════════════════
@router.get("", response_class=HTMLResponse)
@router.get("/", response_class=HTMLResponse)
async def onboard_home(
    request: Request,
    vessel: Optional[int] = Query(None),
    leg_id: Optional[int] = Query(None),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    from app.permissions import can_view
    if not can_view(user, "captain"):
        raise HTTPException(status_code=403)

    now = datetime.now(timezone.utc)
    current_year = now.year

    vessels_result = await db.execute(select(Vessel).where(Vessel.is_active == True).order_by(Vessel.code))
    vessels = vessels_result.scalars().all()

    selected_vessel = vessel or (vessels[0].code if vessels else None)
    vessel_obj = None
    if selected_vessel:
        v_result = await db.execute(select(Vessel).where(Vessel.code == selected_vessel))
        vessel_obj = v_result.scalar_one_or_none()

    legs = []
    current_leg = None
    crew_onboard = []
    cargo_summary = {}
    pax_bookings = []
    sof_events = []
    notifications = []
    last_sof = None
    eta_shifts = []
    attachments = []
    cargo_documents = []

    next_leg = None  # leg export (départ du port)

    if vessel_obj:
        legs_result = await db.execute(
            select(Leg).options(
                selectinload(Leg.departure_port), selectinload(Leg.arrival_port),
            ).where(Leg.vessel_id == vessel_obj.id, Leg.year == current_year)
            .order_by(Leg.sequence)
        )
        legs = legs_result.scalars().all()

        # Select leg: explicit or auto-detect current
        if leg_id:
            current_leg = next((l for l in legs if l.id == leg_id), None)
        if not current_leg:
            for leg in legs:
                if leg.ata and not leg.atd:
                    current_leg = leg
                    break
            if not current_leg and legs:
                # Take next upcoming
                for leg in legs:
                    if not leg.ata:
                        current_leg = leg
                        break
                if not current_leg:
                    current_leg = legs[-1]

        if current_leg:
            # ─── Next leg (export) ───
            next_leg_result = await db.execute(
                select(Leg).options(
                    selectinload(Leg.departure_port), selectinload(Leg.arrival_port),
                ).where(
                    Leg.vessel_id == vessel_obj.id,
                    Leg.year == current_year,
                    Leg.sequence > current_leg.sequence,
                ).order_by(Leg.sequence).limit(1)
            )
            next_leg = next_leg_result.scalar_one_or_none()

            # ─── CREW on board ───
            crew_result = await db.execute(
                select(CrewAssignment).options(selectinload(CrewAssignment.member))
                .where(
                    CrewAssignment.vessel_id == vessel_obj.id,
                    CrewAssignment.status == "active",
                ).order_by(CrewAssignment.member_id)
            )
            crew_onboard = crew_result.scalars().all()

            # ─── CARGO summary ───
            orders_result = await db.execute(
                select(Order).where(Order.leg_id == current_leg.id)
            )
            orders = orders_result.scalars().all()
            order_ids = [o.id for o in orders]

            total_palettes = sum(o.quantity_palettes or 0 for o in orders)
            total_orders = len(orders)

            batches = []
            if order_ids:
                pl_result = await db.execute(
                    select(PackingList).options(selectinload(PackingList.batches))
                    .where(PackingList.order_id.in_(order_ids))
                )
                pls = pl_result.scalars().all()
                for pl in pls:
                    batches.extend(pl.batches)

            total_weight = sum(b.weight_kg or 0 for b in batches)
            total_batch_palettes = sum(b.pallet_quantity or 0 for b in batches)
            goods_types = list(set(b.type_of_goods for b in batches if b.type_of_goods))

            cargo_summary = {
                "orders": total_orders,
                "palettes_ordered": total_palettes,
                "palettes_batches": total_batch_palettes,
                "weight_kg": total_weight,
                "goods_types": goods_types,
                "batches_count": len(batches),
            }

            # ─── SOF events ───
            sof_result = await db.execute(
                select(SofEvent).where(SofEvent.leg_id == current_leg.id)
                .order_by(SofEvent.event_date.asc().nulls_last(), SofEvent.event_time.asc().nulls_last(), SofEvent.id.asc())
            )
            sof_events = sof_result.scalars().all()
            last_sof = sof_events[-1] if sof_events else None

            # ─── Notifications ───
            notif_result = await db.execute(
                select(OnboardNotification).where(OnboardNotification.leg_id == current_leg.id)
                .order_by(OnboardNotification.created_at.desc())
            )
            notifications = notif_result.scalars().all()

            # ─── PASSENGERS for this leg ───
            from app.models.passenger import PassengerBooking, Passenger, PassengerDocument, DOCUMENT_TYPES, CABIN_CONFIG
            pax_result = await db.execute(
                select(PassengerBooking).options(
                    selectinload(PassengerBooking.passengers).selectinload(Passenger.documents),
                )
                .where(
                    PassengerBooking.leg_id == current_leg.id,
                    PassengerBooking.status.notin_(["cancelled"]),
                )
                .order_by(PassengerBooking.cabin_number)
            )
            pax_bookings = pax_result.scalars().all()

            # ─── ETA SHIFTS history for this leg + all vessel legs this year ───
            eta_shifts_result = await db.execute(
                select(ETAShift).where(
                    ETAShift.vessel_id == vessel_obj.id,
                    ETAShift.leg_id.in_([l.id for l in legs]),
                ).order_by(ETAShift.created_at.desc())
            )
            eta_shifts = eta_shifts_result.scalars().all()

            # ─── ATTACHMENTS for this leg ───
            attach_result = await db.execute(
                select(OnboardAttachment).where(OnboardAttachment.leg_id == current_leg.id)
                .order_by(OnboardAttachment.created_at.desc())
            )
            attachments = attach_result.scalars().all()

            # ─── CARGO DOCUMENTS for closure checklist ───
            cargo_docs_result = await db.execute(
                select(CargoDocument).where(CargoDocument.leg_id == current_leg.id)
                .order_by(CargoDocument.created_at)
            )
            cargo_documents = cargo_docs_result.scalars().all()

    # ─── Port timezone for sidebar clock + time inputs ───
    _port_tz = "UTC"
    _port_tz_label = "Port"
    if current_leg:
        # Use arrival port if vessel has arrived, else departure port
        _ref_port = current_leg.arrival_port if current_leg.ata else current_leg.departure_port
        if _ref_port:
            _port_tz = get_port_timezone(_ref_port.country_code, _ref_port.zone_code)
            _port_tz_label = _ref_port.name

    return templates.TemplateResponse("onboard/index.html", {
        "request": request, "user": user,
        "vessels": vessels, "selected_vessel": selected_vessel, "vessel_obj": vessel_obj,
        "legs": legs, "current_leg": current_leg, "next_leg": next_leg,
        "crew_onboard": crew_onboard,
        "cargo_summary": cargo_summary,
        "sof_events": sof_events, "last_sof": last_sof,
        "notifications": notifications,
        "pax_bookings": pax_bookings,
        "eta_shifts": eta_shifts if current_leg else [],
        "eta_shift_reasons": ETA_SHIFT_REASONS,
        "attachments": attachments if current_leg else [],
        "cargo_documents": cargo_documents if current_leg else [],
        "port_agent_attachments": [a for a in (attachments if current_leg else []) if a.category in ("port_agent", "bl_signed", "letter_protest")],
        "attachment_categories": ATTACHMENT_CATEGORIES,
        "sof_event_types": SOF_EVENT_TYPES,
        "cargo_doc_types": CARGO_DOC_TYPES,
        "tz_choices": TIMEZONE_CHOICES,
        "port_timezone": _port_tz,
        "port_tz_label": _port_tz_label,
        "port_tz_offset": utc_offset_label(_port_tz),
        "current_year": current_year,
        "active_module": "captain",
        "lang": user.language or "fr",
    })


# ═══════════════════════════════════════════════════════════════
#  SOF EVENTS
# ═══════════════════════════════════════════════════════════════
@router.post("/sof/add", response_class=HTMLResponse)
async def sof_add_event(
    request: Request,
    leg_id: int = Form(...),
    event_type: str = Form(...),
    event_label: str = Form(""),
    event_date: str = Form(""),
    event_time: str = Form(""),
    event_time_tz: str = Form("UTC"),
    remarks: str = Form(""),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    # Resolve 'port_local' timezone to actual IANA name
    resolved_tz = event_time_tz
    if event_time_tz == "port_local":
        leg = await db.get(Leg, leg_id)
        if leg:
            from app.models.port import Port
            dep = await db.get(Port, leg.departure_port_id) if leg.departure_port_id else None
            arr = await db.get(Port, leg.arrival_port_id) if leg.arrival_port_id else None
            ref_port = arr if leg.ata else dep
            if ref_port:
                resolved_tz = get_port_timezone(ref_port.country_code, ref_port.zone_code)

    # Find label from type if not custom
    label = event_label.strip()
    if not label:
        label = next((l for c, l in SOF_EVENT_TYPES if c == event_type), event_type)

    evt = SofEvent(
        leg_id=leg_id,
        event_type=event_type,
        event_label=label,
        event_date=datetime.strptime(event_date, "%Y-%m-%d").date() if event_date else date.today(),
        event_time=event_time or None,
        event_time_tz=resolved_tz,
        remarks=remarks or None,
        created_by=user.full_name,
    )
    db.add(evt)
    await db.flush()

    # Notification EOSP / SOSP
    if event_type in ("EOSP", "SOSP"):
        from app.models.notification import Notification
        leg = await db.get(Leg, leg_id)
        vessel_obj = await db.get(Vessel, leg.vessel_id) if leg else None
        vessel_name = vessel_obj.name if vessel_obj else ""
        db.add(Notification(
            type=event_type.lower(),
            title=f"{event_type} — {leg.leg_code if leg else ''}",
            detail=f"{vessel_name} · {label}",
            link=f"/onboard?vessel={vessel_obj.code if vessel_obj else ''}&leg_id={leg_id}#sof",
            leg_id=leg_id,
        ))
        await db.flush()

    # Find vessel for redirect
    leg = await db.get(Leg, leg_id)
    vessel_obj = await db.get(Vessel, leg.vessel_id) if leg else None
    vc = vessel_obj.code if vessel_obj else ""
    return RedirectResponse(url=f"/onboard?vessel={vc}&leg_id={leg_id}#sof", status_code=303)


@router.post("/sof/{event_id}/edit", response_class=HTMLResponse)
async def sof_edit_event(
    event_id: int, request: Request,
    event_label: str = Form(""),
    event_date: str = Form(""),
    event_time: str = Form(""),
    event_time_tz: str = Form("UTC"),
    remarks: str = Form(""),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    evt = await db.get(SofEvent, event_id)
    if not evt:
        raise HTTPException(404)
    if event_label.strip():
        evt.event_label = event_label.strip()
    if event_date:
        evt.event_date = datetime.strptime(event_date, "%Y-%m-%d").date()
    evt.event_time = event_time or evt.event_time
    # Resolve port_local timezone
    resolved_tz = event_time_tz
    if event_time_tz == "port_local":
        leg = await db.get(Leg, evt.leg_id)
        if leg:
            from app.models.port import Port
            dep = await db.get(Port, leg.departure_port_id) if leg.departure_port_id else None
            arr = await db.get(Port, leg.arrival_port_id) if leg.arrival_port_id else None
            ref_port = arr if leg.ata else dep
            if ref_port:
                resolved_tz = get_port_timezone(ref_port.country_code, ref_port.zone_code)
    evt.event_time_tz = resolved_tz
    evt.remarks = remarks if remarks is not None else evt.remarks
    await db.flush()

    leg = await db.get(Leg, evt.leg_id)
    vessel_obj = await db.get(Vessel, leg.vessel_id) if leg else None
    vc = vessel_obj.code if vessel_obj else ""
    return RedirectResponse(url=f"/onboard?vessel={vc}&leg_id={evt.leg_id}#sof", status_code=303)


@router.delete("/sof/{event_id}", response_class=HTMLResponse)
async def sof_delete_event(
    event_id: int, request: Request,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    evt = await db.get(SofEvent, event_id)
    if evt:
        await db.delete(evt)
        await db.flush()
    return HTMLResponse(content="", status_code=200)


# ═══════════════════════════════════════════════════════════════
#  NOTIFICATIONS
# ═══════════════════════════════════════════════════════════════
@router.post("/notifications/{notif_id}/dismiss", response_class=HTMLResponse)
async def dismiss_notification(
    notif_id: int, request: Request,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    notif = await db.get(OnboardNotification, notif_id)
    if notif:
        notif.is_read = True
        await db.flush()
    return HTMLResponse(content="", status_code=200)


@router.post("/notifications/dismiss-all", response_class=HTMLResponse)
async def dismiss_all_notifications(
    request: Request,
    leg_id: int = Form(...),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    from sqlalchemy import update
    await db.execute(
        update(OnboardNotification)
        .where(OnboardNotification.leg_id == leg_id, OnboardNotification.is_read == False)
        .values(is_read=True)
    )
    await db.flush()
    leg = await db.get(Leg, leg_id)
    vessel_obj = await db.get(Vessel, leg.vessel_id) if leg else None
    vc = vessel_obj.code if vessel_obj else ""
    return RedirectResponse(url=f"/onboard?vessel={vc}&leg_id={leg_id}", status_code=303)


# ═══════════════════════════════════════════════════════════════
#  ATA / ATD — Captain records actual arrival and departure
# ═══════════════════════════════════════════════════════════════

@router.post("/set-ata", response_class=HTMLResponse)
async def set_actual_arrival(
    request: Request,
    leg_id: int = Form(...),
    ata_date: str = Form(...),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Captain records actual time of arrival (ATA)."""
    from app.models.notification import Notification
    from app.routers.planning_router import resequence_and_recalc

    leg_result = await db.execute(
        select(Leg).options(
            selectinload(Leg.vessel), selectinload(Leg.departure_port), selectinload(Leg.arrival_port),
        ).where(Leg.id == leg_id)
    )
    leg = leg_result.scalar_one_or_none()
    if not leg:
        raise HTTPException(404, "Leg non trouve")

    ata_dt = None
    if ata_date and ata_date.strip():
        try:
            ata_dt = datetime.fromisoformat(ata_date)
        except ValueError:
            raise HTTPException(400, "Format de date invalide")

    if ata_dt and leg.etd and ata_dt < leg.etd:
        raise HTTPException(400, "L'ATA ne peut pas etre avant l'ETD")

    old_ata = leg.ata
    leg.ata = ata_dt
    await db.flush()

    # Cascade to downstream legs
    await resequence_and_recalc(db, leg.vessel_id, leg.year)

    # Create SOF event
    sof = SofEvent(
        leg_id=leg_id,
        event_type="EOSP",
        event_label=f"Arrivee confirmee — ATA {ata_dt.strftime('%d/%m/%Y %H:%M') if ata_dt else 'annulee'}",
        event_date=ata_dt.date() if ata_dt else None,
        event_time=ata_dt.strftime("%H:%M") if ata_dt else None,
        remarks=f"ATA enregistree par {user.full_name}" + (f" (ancien: {old_ata.strftime('%d/%m %H:%M')})" if old_ata else ""),
        created_by=user.full_name,
    )
    db.add(sof)

    # Notify
    vessel_name = leg.vessel.name if leg.vessel else ""
    port_name = leg.arrival_port.name if leg.arrival_port else leg.arrival_port_locode
    db.add(Notification(
        type="ata_recorded",
        title=f"ATA {vessel_name} — {port_name}",
        detail=f"Arrivee confirmee le {ata_dt.strftime('%d/%m/%Y %H:%M') if ata_dt else '—'} par {user.full_name}",
        link=f"/planning?vessel={leg.vessel.code if leg.vessel else ''}&year={leg.year}",
        leg_id=leg_id,
    ))
    await db.flush()

    vc = leg.vessel.code if leg.vessel else ""
    return RedirectResponse(url=f"/onboard?vessel={vc}&leg_id={leg_id}", status_code=303)


@router.post("/set-atd", response_class=HTMLResponse)
async def set_actual_departure(
    request: Request,
    leg_id: int = Form(...),
    atd_date: str = Form(...),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Captain records actual time of departure (ATD)."""
    from app.models.notification import Notification
    from app.routers.planning_router import resequence_and_recalc

    leg_result = await db.execute(
        select(Leg).options(
            selectinload(Leg.vessel), selectinload(Leg.departure_port), selectinload(Leg.arrival_port),
        ).where(Leg.id == leg_id)
    )
    leg = leg_result.scalar_one_or_none()
    if not leg:
        raise HTTPException(404, "Leg non trouve")

    atd_dt = None
    if atd_date and atd_date.strip():
        try:
            atd_dt = datetime.fromisoformat(atd_date)
        except ValueError:
            raise HTTPException(400, "Format de date invalide")

    if atd_dt and leg.ata and atd_dt < leg.ata:
        raise HTTPException(400, "L'ATD ne peut pas etre avant l'ATA")

    old_atd = leg.atd
    leg.atd = atd_dt
    await db.flush()

    # Cascade to downstream legs
    await resequence_and_recalc(db, leg.vessel_id, leg.year)

    # Create SOF event
    sof = SofEvent(
        leg_id=leg_id,
        event_type="SOSP",
        event_label=f"Depart confirme — ATD {atd_dt.strftime('%d/%m/%Y %H:%M') if atd_dt else 'annule'}",
        event_date=atd_dt.date() if atd_dt else None,
        event_time=atd_dt.strftime("%H:%M") if atd_dt else None,
        remarks=f"ATD enregistre par {user.full_name}" + (f" (ancien: {old_atd.strftime('%d/%m %H:%M')})" if old_atd else ""),
        created_by=user.full_name,
    )
    db.add(sof)

    # Notify
    vessel_name = leg.vessel.name if leg.vessel else ""
    port_name = leg.departure_port.name if leg.departure_port else leg.departure_port_locode
    db.add(Notification(
        type="atd_recorded",
        title=f"ATD {vessel_name} — {port_name}",
        detail=f"Depart confirme le {atd_dt.strftime('%d/%m/%Y %H:%M') if atd_dt else '—'} par {user.full_name}",
        link=f"/planning?vessel={leg.vessel.code if leg.vessel else ''}&year={leg.year}",
        leg_id=leg_id,
    ))
    await db.flush()

    vc = leg.vessel.code if leg.vessel else ""
    return RedirectResponse(url=f"/onboard?vessel={vc}&leg_id={leg_id}", status_code=303)


# ═══════════════════════════════════════════════════════════════
#  ETA SHIFT — Captain declares ETA/ETD change during navigation
# ═══════════════════════════════════════════════════════════════
@router.post("/eta-shift", response_class=HTMLResponse)
async def eta_shift_declare(
    request: Request,
    leg_id: int = Form(...),
    new_eta: str = Form(""),
    reason: str = Form(...),
    justification: str = Form(...),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Captain declares ETA shift for current navigation leg.
    Cascades to all subsequent legs and notifies the company."""
    from app.models.notification import Notification
    from app.routers.planning_router import resequence_and_recalc
    from datetime import timedelta

    leg_result = await db.execute(
        select(Leg).options(
            selectinload(Leg.vessel), selectinload(Leg.departure_port), selectinload(Leg.arrival_port),
        ).where(Leg.id == leg_id)
    )
    leg = leg_result.scalar_one_or_none()
    if not leg:
        raise HTTPException(404, "Leg non trouvé")

    # Parse new ETA
    new_eta_dt = None
    if new_eta and new_eta.strip():
        try:
            new_eta_dt = datetime.fromisoformat(new_eta)
        except ValueError:
            raise HTTPException(400, "Format de date invalide")

    if not new_eta_dt:
        raise HTTPException(400, "Nouvelle ETA requise")

    # Validate justification
    if not justification or not justification.strip():
        raise HTTPException(400, "La justification est obligatoire")

    old_eta = leg.eta
    shift_hours = round((new_eta_dt - old_eta).total_seconds() / 3600, 1) if old_eta else 0

    # Record the shift in history
    shift_record = ETAShift(
        leg_id=leg.id,
        vessel_id=leg.vessel_id,
        field_changed="eta",
        old_value=old_eta,
        new_value=new_eta_dt,
        shift_hours=shift_hours,
        reason=reason,
        justification=justification.strip(),
        created_by=user.full_name,
    )
    db.add(shift_record)

    # Apply the new ETA
    leg.eta = new_eta_dt
    await db.flush()

    # Cascade recalculation to all subsequent legs
    await resequence_and_recalc(db, leg.vessel_id, leg.year)

    # Count affected downstream legs
    downstream_result = await db.execute(
        select(func.count(Leg.id)).where(
            Leg.vessel_id == leg.vessel_id,
            Leg.year == leg.year,
            Leg.sequence > leg.sequence,
        )
    )
    legs_affected = downstream_result.scalar() or 0
    shift_record.legs_affected = legs_affected
    await db.flush()

    # Create SOF event
    direction = "retard" if shift_hours > 0 else "avance"
    sof_evt = SofEvent(
        leg_id=leg.id,
        event_type="CUSTOM",
        event_label=f"ETA modifiée : {direction} de {abs(shift_hours):.1f}h",
        event_date=datetime.now().date(),
        event_time=datetime.now().strftime("%H:%M"),
        remarks=f"Raison: {reason} — {justification.strip()}",
        created_by=user.full_name,
    )
    db.add(sof_evt)

    # Create company-wide notification
    vessel_name = leg.vessel.name if leg.vessel else ""
    sign = "+" if shift_hours > 0 else ""
    reason_label = next((l for c, l in ETA_SHIFT_REASONS if c == reason), reason)
    detail_msg = (
        f"{vessel_name} · {leg.leg_code} · "
        f"ETA {old_eta.strftime('%d/%m %H:%M') if old_eta else '—'} → {new_eta_dt.strftime('%d/%m %H:%M')} "
        f"({sign}{shift_hours:.1f}h) · {reason_label}"
    )
    if legs_affected > 0:
        detail_msg += f" · {legs_affected} leg(s) suivant(s) recalculé(s)"

    db.add(Notification(
        type="eta_shift",
        title=f"ETA modifiée — {leg.leg_code} ({sign}{shift_hours:.1f}h)",
        detail=detail_msg,
        link=f"/onboard?vessel={leg.vessel.code if leg.vessel else ''}&leg_id={leg.id}#eta-shifts",
        leg_id=leg.id,
    ))

    # Onboard notification too
    db.add(OnboardNotification(
        leg_id=leg.id,
        category="escale",
        title=f"ETA modifiée : {sign}{shift_hours:.1f}h",
        detail=f"{reason_label} — {justification.strip()}",
    ))

    await db.flush()

    vc = leg.vessel.code if leg.vessel else ""
    return RedirectResponse(url=f"/onboard?vessel={vc}&leg_id={leg_id}#eta-shifts", status_code=303)


# ═══════════════════════════════════════════════════════════════
#  ATTACHMENTS (file/photo upload)
# ═══════════════════════════════════════════════════════════════
UPLOAD_DIR = "/app/uploads/onboard"
ALLOWED_EXTENSIONS = {".pdf", ".jpg", ".jpeg", ".png", ".gif", ".doc", ".docx", ".xls", ".xlsx", ".csv", ".txt", ".heic", ".webp"}
MAX_FILE_SIZE = 20 * 1024 * 1024  # 20 MB


@router.post("/attachments/upload", response_class=HTMLResponse)
async def attachment_upload(
    request: Request,
    leg_id: int = Form(...),
    category: str = Form("document"),
    title: str = Form(""),
    description: str = Form(""),
    file: UploadFile = File(...),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Upload a file or photo attachment to a leg."""
    import os
    import uuid

    leg = await db.get(Leg, leg_id)
    if not leg:
        raise HTTPException(404)

    # Check leg not locked
    if leg.status == "completed":
        raise HTTPException(400, "L'escale est clôturée, impossible d'ajouter des fichiers")

    # Validate extension
    ext = os.path.splitext(file.filename)[1].lower() if file.filename else ""
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(400, f"Type de fichier non autorisé : {ext}")

    # Read and check size
    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(400, "Fichier trop volumineux (max 20 Mo)")

    # Validate file content matches extension (magic bytes)
    from app.utils.file_validation import validate_file_content
    if not validate_file_content(content, ext):
        raise HTTPException(400, "Le contenu du fichier ne correspond pas à son extension")

    # Create upload dir
    os.makedirs(UPLOAD_DIR, exist_ok=True)

    # Safe filename
    safe_name = f"{leg_id}_{uuid.uuid4().hex[:8]}{ext}"
    file_path = os.path.join(UPLOAD_DIR, safe_name)

    with open(file_path, "wb") as f:
        f.write(content)

    att = OnboardAttachment(
        leg_id=leg_id,
        category=category,
        title=title.strip() or file.filename or "Sans titre",
        filename=file.filename or safe_name,
        file_path=file_path,
        file_size=len(content),
        mime_type=file.content_type,
        description=description.strip() or None,
        uploaded_by=user.full_name,
    )
    db.add(att)
    await db.flush()

    vessel_obj = await db.get(Vessel, leg.vessel_id)
    vc = vessel_obj.code if vessel_obj else ""
    return RedirectResponse(url=f"/onboard?vessel={vc}&leg_id={leg_id}#attachments", status_code=303)


@router.get("/attachments/{att_id}/download")
async def attachment_download(
    att_id: int, request: Request,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Download an attachment file."""
    import os
    att = await db.get(OnboardAttachment, att_id)
    if not att or not os.path.isfile(att.file_path):
        raise HTTPException(404)

    with open(att.file_path, "rb") as f:
        content = f.read()

    return Response(
        content=content,
        media_type=att.mime_type or "application/octet-stream",
        headers={"Content-Disposition": f'attachment; filename="{att.filename}"'},
    )


@router.delete("/attachments/{att_id}", response_class=HTMLResponse)
async def attachment_delete(
    att_id: int, request: Request,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Delete an attachment."""
    import os
    att = await db.get(OnboardAttachment, att_id)
    if att:
        # Check leg not locked
        leg = await db.get(Leg, att.leg_id)
        if leg and leg.status == "completed":
            raise HTTPException(400, "L'escale est clôturée")
        if os.path.isfile(att.file_path):
            os.remove(att.file_path)
        await db.delete(att)
        await db.flush()
    return HTMLResponse(content="", status_code=200)


# ═══════════════════════════════════════════════════════════════
#  DOCUMENT ATTACHMENTS (photos/files linked to cargo documents)
# ═══════════════════════════════════════════════════════════════
DOC_UPLOAD_DIR = "/app/uploads/onboard/documents"


@router.post("/doc/{doc_id}/attachments/upload", response_class=HTMLResponse)
async def doc_attachment_upload(
    doc_id: int, request: Request,
    file: UploadFile = File(...),
    title: str = Form(""),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Upload a file or photo to a cargo document."""
    import os
    import uuid

    doc = await db.get(CargoDocument, doc_id)
    if not doc:
        raise HTTPException(404)

    leg = await db.get(Leg, doc.leg_id)
    if leg and leg.status == "completed":
        raise HTTPException(400, "L'escale est clôturée, impossible d'ajouter des fichiers")

    ext = os.path.splitext(file.filename)[1].lower() if file.filename else ""
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(400, f"Type de fichier non autorisé : {ext}")

    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(400, "Fichier trop volumineux (max 20 Mo)")

    # Validate file content matches extension (magic bytes)
    from app.utils.file_validation import validate_file_content
    if not validate_file_content(content, ext):
        raise HTTPException(400, "Le contenu du fichier ne correspond pas à son extension")

    os.makedirs(DOC_UPLOAD_DIR, exist_ok=True)
    safe_name = f"doc{doc_id}_{uuid.uuid4().hex[:8]}{ext}"
    file_path = os.path.join(DOC_UPLOAD_DIR, safe_name)

    with open(file_path, "wb") as f:
        f.write(content)

    att = CargoDocumentAttachment(
        document_id=doc_id,
        leg_id=doc.leg_id,
        title=title.strip() or file.filename or "Sans titre",
        filename=file.filename or safe_name,
        file_path=file_path,
        file_size=len(content),
        mime_type=file.content_type,
        uploaded_by=user.full_name,
    )
    db.add(att)
    await db.flush()

    return RedirectResponse(
        url=f"/onboard/doc/{doc.leg_id}/{doc.doc_type}?doc_id={doc_id}#doc-attachments",
        status_code=303,
    )


@router.get("/doc/attachments/{att_id}/download")
async def doc_attachment_download(
    att_id: int, request: Request,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Download a document attachment."""
    import os
    att = await db.get(CargoDocumentAttachment, att_id)
    if not att or not os.path.isfile(att.file_path):
        raise HTTPException(404)

    with open(att.file_path, "rb") as f:
        content = f.read()

    return Response(
        content=content,
        media_type=att.mime_type or "application/octet-stream",
        headers={"Content-Disposition": f'attachment; filename="{att.filename}"'},
    )


@router.delete("/doc/attachments/{att_id}", response_class=HTMLResponse)
async def doc_attachment_delete(
    att_id: int, request: Request,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Delete a document attachment."""
    import os
    att = await db.get(CargoDocumentAttachment, att_id)
    if att:
        leg = await db.get(Leg, att.leg_id)
        if leg and leg.status == "completed":
            raise HTTPException(400, "L'escale est clôturée")
        if os.path.isfile(att.file_path):
            os.remove(att.file_path)
        await db.delete(att)
        await db.flush()
    return HTMLResponse(content="", status_code=200)


# ═══════════════════════════════════════════════════════════════
#  SOF EXPORT (Excel + PDF)
# ═══════════════════════════════════════════════════════════════
@router.get("/sof/{leg_id}/excel", response_class=StreamingResponse)
async def sof_export_excel(
    leg_id: int, request: Request,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    from io import BytesIO
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

    leg = await db.get(Leg, leg_id)
    if not leg:
        raise HTTPException(404)
    vessel = await db.get(Vessel, leg.vessel_id)

    sof_result = await db.execute(
        select(SofEvent).where(SofEvent.leg_id == leg_id)
        .order_by(SofEvent.event_date.asc().nulls_last(), SofEvent.event_time.asc().nulls_last(), SofEvent.id.asc())
    )
    events = sof_result.scalars().all()

    wb = Workbook()
    ws = wb.active
    ws.title = "Statement of Facts"

    # Header
    hdr_fill = PatternFill(start_color="095561", end_color="095561", fill_type="solid")
    hdr_font = Font(bold=True, color="FFFFFF", size=12)
    ws.merge_cells("A1:E1")
    ws["A1"] = "STATEMENT OF FACTS"
    ws["A1"].font = hdr_font
    ws["A1"].fill = hdr_fill
    ws["A1"].alignment = Alignment(horizontal="center")
    for c in ["B1", "C1", "D1", "E1"]:
        ws[c].fill = hdr_fill

    ws["A3"] = "Vessel:"
    ws["B3"] = vessel.name if vessel else ""
    ws["D3"] = "Voyage:"
    ws["E3"] = leg.leg_code

    # Table header
    headers = ["Event", "Date", "Time (LT)", "Remarks", "Recorded by"]
    row_font = Font(bold=True, color="FFFFFF", size=10)
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=5, column=col, value=h)
        cell.font = row_font
        cell.fill = PatternFill(start_color="095561", end_color="095561", fill_type="solid")

    for i, evt in enumerate(events, 6):
        ws.cell(row=i, column=1, value=evt.event_label)
        ws.cell(row=i, column=2, value=evt.event_date.strftime("%d/%m/%Y") if evt.event_date else "")
        ws.cell(row=i, column=3, value=evt.event_time or "")
        ws.cell(row=i, column=4, value=evt.remarks or "")
        ws.cell(row=i, column=5, value=evt.created_by or "")

    ws.column_dimensions["A"].width = 50
    ws.column_dimensions["B"].width = 14
    ws.column_dimensions["C"].width = 12
    ws.column_dimensions["D"].width = 40
    ws.column_dimensions["E"].width = 20

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)

    fname = f"SOF_{leg.leg_code}_{date.today().strftime('%Y%m%d')}.xlsx"
    return StreamingResponse(
        buf, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={fname}"},
    )


@router.get("/sof/{leg_id}/pdf", response_class=StreamingResponse)
async def sof_export_pdf(
    leg_id: int, request: Request,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Export SOF as PDF using reportlab."""
    from io import BytesIO
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

    leg_result = await db.execute(
        select(Leg).options(selectinload(Leg.vessel), selectinload(Leg.departure_port), selectinload(Leg.arrival_port))
        .where(Leg.id == leg_id)
    )
    leg = leg_result.scalar_one_or_none()
    if not leg:
        raise HTTPException(404)

    sof_result = await db.execute(
        select(SofEvent).where(SofEvent.leg_id == leg_id)
        .order_by(SofEvent.event_date.asc().nulls_last(), SofEvent.event_time.asc().nulls_last(), SofEvent.id.asc())
    )
    events = sof_result.scalars().all()

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=20*mm, bottomMargin=15*mm, leftMargin=15*mm, rightMargin=15*mm)
    styles = getSampleStyleSheet()
    elements = []

    # Title
    title_style = ParagraphStyle("SOFTitle", parent=styles["Heading1"], fontSize=16, textColor=colors.HexColor("#095561"), spaceAfter=6)
    elements.append(Paragraph("STATEMENT OF FACTS", title_style))

    # Header info
    info_style = ParagraphStyle("Info", parent=styles["Normal"], fontSize=10, textColor=colors.HexColor("#555555"), spaceAfter=2)
    elements.append(Paragraph(f"<b>Vessel:</b> {leg.vessel.name if leg.vessel else ''} &nbsp;&nbsp; <b>Voyage:</b> {leg.leg_code}", info_style))
    elements.append(Paragraph(f"<b>From:</b> {leg.departure_port.name if leg.departure_port else ''} &nbsp;&nbsp; <b>To:</b> {leg.arrival_port.name if leg.arrival_port else ''}", info_style))
    elements.append(Spacer(1, 10*mm))

    # Table
    hdr_color = colors.HexColor("#095561")
    data = [["Event", "Date", "Time (LT)", "Remarks"]]
    for evt in events:
        data.append([
            evt.event_label,
            evt.event_date.strftime("%d/%m/%Y") if evt.event_date else "",
            evt.event_time or "",
            evt.remarks or "",
        ])

    if len(data) > 1:
        col_widths = [220, 70, 60, 180]
        t = Table(data, colWidths=col_widths, repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), hdr_color),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, 0), 9),
            ("FONTSIZE", (0, 1), (-1, -1), 8),
            ("ALIGN", (0, 0), (-1, -1), "LEFT"),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#dddddd")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f9fafb")]),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        elements.append(t)
    else:
        elements.append(Paragraph("No SOF events recorded.", info_style))

    # Footer
    elements.append(Spacer(1, 10*mm))
    footer_style = ParagraphStyle("Footer", parent=styles["Normal"], fontSize=8, textColor=colors.HexColor("#aaaaaa"))
    elements.append(Paragraph(f"Generated on {date.today().strftime('%d/%m/%Y')} — TOWT Operations Platform", footer_style))

    doc.build(elements)
    buf.seek(0)

    fname = f"SOF_{leg.leg_code}_{date.today().strftime('%Y%m%d')}.pdf"
    return StreamingResponse(
        buf, media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={fname}"},
    )


# ═══════════════════════════════════════════════════════════════
#  CARGO DOCUMENT EXPORT (Word + PDF) + SOF historisation
# ═══════════════════════════════════════════════════════════════
def _build_doc_paragraphs(doc_data: dict, doc_type: str, prefill: dict) -> list:
    """Build a list of (label, value) pairs for document rendering."""
    # Common header fields
    rows = [
        ("Vessel", doc_data.get("vessel_name", prefill.get("vessel_name", ""))),
        ("Voyage No", doc_data.get("voyage_no", prefill.get("voyage_no", ""))),
    ]

    if doc_type == "NOR":
        rows += [
            ("To", doc_data.get("to_charterer", "")),
            ("Port", doc_data.get("port", "")),
            ("Date of Notice", doc_data.get("notice_date", "")),
            ("Time of Notice (LT)", doc_data.get("notice_time", "")),
            ("Cargo description", doc_data.get("cargo_desc", "")),
            ("Position", doc_data.get("position", "")),
            ("Remarks", doc_data.get("remarks", "")),
            ("Master", doc_data.get("master_name", "")),
        ]
    elif doc_type == "NOR_RT":
        rows += [
            ("To", doc_data.get("to_charterer", "")),
            ("Port", doc_data.get("port", "")),
            ("Reason for re-tendering", doc_data.get("reason", "")),
            ("Date", doc_data.get("notice_date", "")),
            ("Time (LT)", doc_data.get("notice_time", "")),
            ("Master", doc_data.get("master_name", "")),
        ]
    elif doc_type == "HOLDS_CERT":
        rows += [
            ("To", doc_data.get("to", "")),
            ("Port", doc_data.get("port", "")),
            ("Cargo to be loaded", doc_data.get("cargo", "")),
            ("Date of inspection", doc_data.get("inspection_date", "")),
            ("Holds inspected", doc_data.get("holds_list", "")),
            ("Result / Observations", doc_data.get("observations", "")),
            ("Chief Officer / Master", doc_data.get("officer_name", "")),
            ("Surveyor / Terminal", doc_data.get("surveyor", "")),
        ]
    elif doc_type in ("KEY_MEETING", "PRE_MEETING"):
        rows += [
            ("Port", doc_data.get("port", "")),
            ("Date", doc_data.get("meeting_date", "")),
            ("Attendees", doc_data.get("attendees", doc_data.get("terminal", ""))),
            ("Content", doc_data.get("key_points", doc_data.get("plan", ""))),
            ("Actions", doc_data.get("actions", doc_data.get("emergency", ""))),
        ]
    elif doc_type.startswith("LOP"):
        rows += [
            ("To", doc_data.get("to", "")),
            ("Port", doc_data.get("port", "")),
            ("Date", doc_data.get("lop_date", "")),
            ("Time", doc_data.get("lop_time", "")),
            ("Subject", doc_data.get("subject", "")),
            ("Details", doc_data.get("details", "")),
            ("Reserve", doc_data.get("reserve", "")),
            ("Master", doc_data.get("master_name", "")),
            ("Countersigned", doc_data.get("countersigned", "")),
        ]
    elif doc_type == "MATES_RECEIPT":
        rows += [
            ("Port of loading", doc_data.get("port_loading", "")),
            ("Date", doc_data.get("receipt_date", "")),
            ("Shipper", doc_data.get("shipper", "")),
            ("Cargo description", doc_data.get("cargo_desc", "")),
            ("Number of packages", doc_data.get("packages", "")),
            ("Gross weight (kg)", doc_data.get("weight", "")),
            ("Condition / Remarks", doc_data.get("condition", "")),
            ("Chief Officer", doc_data.get("officer_name", "")),
        ]
    else:
        rows += [
            ("Port", doc_data.get("port", "")),
            ("Date", doc_data.get("doc_date", "")),
            ("Content", doc_data.get("content", "")),
        ]

    return [(k, v) for k, v in rows if v]


@router.get("/doc/{doc_id}/export/word")
async def cargo_doc_export_word(
    doc_id: int, request: Request,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Export a cargo document as Word (.docx)."""
    from io import BytesIO
    from fastapi.responses import Response
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    cargo_doc = await db.get(CargoDocument, doc_id)
    if not cargo_doc:
        raise HTTPException(404)

    leg_result = await db.execute(
        select(Leg).options(selectinload(Leg.vessel), selectinload(Leg.departure_port), selectinload(Leg.arrival_port))
        .where(Leg.id == cargo_doc.leg_id)
    )
    leg = leg_result.scalar_one_or_none()

    doc_data = json.loads(cargo_doc.data_json) if cargo_doc.data_json else {}
    prefill = {
        "vessel_name": leg.vessel.name if leg and leg.vessel else "",
        "voyage_no": leg.leg_code if leg else "",
    }
    rows = _build_doc_paragraphs(doc_data, cargo_doc.doc_type, prefill)
    doc_label = next((l for c, l in CARGO_DOC_TYPES if c == cargo_doc.doc_type), cargo_doc.doc_type)

    d = DocxDocument()

    title = d.add_heading(doc_label.upper(), level=1)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x09, 0x55, 0x61)

    d.add_paragraph(f"TOWT — {leg.vessel.name if leg and leg.vessel else ''} — {leg.leg_code if leg else ''}")

    if rows:
        table = d.add_table(rows=len(rows), cols=2)
        table.style = "Table Grid"
        for i, (label, value) in enumerate(rows):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = str(value) if value else ""
            for p in table.cell(i, 0).paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(10)

    d.add_paragraph("")
    sig = d.add_paragraph()
    sig.add_run("\n\nMaster signature: ____________________________").font.size = Pt(10)
    sig.add_run("\n\nDate: ____________________________").font.size = Pt(10)

    footer_p = d.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run(f"Generated {date.today().strftime('%d/%m/%Y')} — TOWT Operations Platform")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    buf = BytesIO()
    d.save(buf)
    content = buf.getvalue()

    # Historize in SOF - do this BEFORE returning response
    sof_evt = SofEvent(
        leg_id=cargo_doc.leg_id,
        event_type="CUSTOM",
        event_label=f"Document generated: {doc_label} (Word)",
        event_date=date.today(),
        event_time=datetime.now().strftime("%H:%M"),
        remarks=f"Exported by {user.full_name}",
        created_by=user.full_name,
    )
    db.add(sof_evt)
    await db.flush()

    fname = f"{cargo_doc.doc_type}_{leg.leg_code}_{date.today().strftime('%Y%m%d')}.docx"
    return Response(
        content=content,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={fname}"},
    )


@router.get("/doc/{doc_id}/export/pdf")
async def cargo_doc_export_pdf(
    doc_id: int, request: Request,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Export a cargo document as PDF."""
    from io import BytesIO
    from fastapi.responses import Response
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

    cargo_doc = await db.get(CargoDocument, doc_id)
    if not cargo_doc:
        raise HTTPException(404)

    leg_result = await db.execute(
        select(Leg).options(selectinload(Leg.vessel), selectinload(Leg.departure_port), selectinload(Leg.arrival_port))
        .where(Leg.id == cargo_doc.leg_id)
    )
    leg = leg_result.scalar_one_or_none()

    doc_data = json.loads(cargo_doc.data_json) if cargo_doc.data_json else {}
    prefill = {
        "vessel_name": leg.vessel.name if leg and leg.vessel else "",
        "voyage_no": leg.leg_code if leg else "",
    }
    rows = _build_doc_paragraphs(doc_data, cargo_doc.doc_type, prefill)
    doc_label = next((l for c, l in CARGO_DOC_TYPES if c == cargo_doc.doc_type), cargo_doc.doc_type)

    buf = BytesIO()
    pdf_doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=20*mm, bottomMargin=15*mm, leftMargin=15*mm, rightMargin=15*mm)
    styles = getSampleStyleSheet()
    elements = []

    hdr_color = colors.HexColor("#095561")
    title_style = ParagraphStyle("DocTitle", parent=styles["Heading1"], fontSize=16, textColor=hdr_color, spaceAfter=4)
    elements.append(Paragraph(doc_label.upper(), title_style))

    sub_style = ParagraphStyle("Sub", parent=styles["Normal"], fontSize=10, textColor=colors.HexColor("#555555"), spaceAfter=2)
    elements.append(Paragraph(f"TOWT — {leg.vessel.name if leg and leg.vessel else ''} — {leg.leg_code if leg else ''}", sub_style))
    elements.append(Spacer(1, 8*mm))

    data = [[k, str(v) if v else ""] for k, v in rows]
    if data:
        t = Table(data, colWidths=[130, 400])
        t.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#dddddd")),
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f0f7f8")),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]))
        elements.append(t)

    elements.append(Spacer(1, 15*mm))
    sig_style = ParagraphStyle("Sig", parent=styles["Normal"], fontSize=10, spaceAfter=8)
    elements.append(Paragraph("Master signature: ____________________________", sig_style))
    elements.append(Paragraph("Date: ____________________________", sig_style))

    elements.append(Spacer(1, 10*mm))
    footer_style = ParagraphStyle("Footer", parent=styles["Normal"], fontSize=8, textColor=colors.HexColor("#aaaaaa"))
    elements.append(Paragraph(f"Generated on {date.today().strftime('%d/%m/%Y')} — TOWT Operations Platform", footer_style))

    pdf_doc.build(elements)
    content = buf.getvalue()

    # Historize in SOF - BEFORE returning
    sof_evt = SofEvent(
        leg_id=cargo_doc.leg_id,
        event_type="CUSTOM",
        event_label=f"Document generated: {doc_label} (PDF)",
        event_date=date.today(),
        event_time=datetime.now().strftime("%H:%M"),
        remarks=f"Exported by {user.full_name}",
        created_by=user.full_name,
    )
    db.add(sof_evt)
    await db.flush()

    fname = f"{cargo_doc.doc_type}_{leg.leg_code}_{date.today().strftime('%Y%m%d')}.pdf"
    return Response(
        content=content,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={fname}"},
    )


# ═══════════════════════════════════════════════════════════════
#  CARGO DOCUMENT GENERATION (form pages)
# ═══════════════════════════════════════════════════════════════
@router.get("/doc/{leg_id}/{doc_type}", response_class=HTMLResponse)
async def cargo_doc_form(
    leg_id: int, doc_type: str, request: Request,
    doc_id: Optional[int] = Query(None),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Show form page for filling a cargo document."""
    leg_result = await db.execute(
        select(Leg).options(
            selectinload(Leg.vessel), selectinload(Leg.departure_port), selectinload(Leg.arrival_port),
        ).where(Leg.id == leg_id)
    )
    leg = leg_result.scalar_one_or_none()
    if not leg:
        raise HTTPException(404)

    # Load existing doc if editing
    doc = None
    doc_data = {}
    if doc_id:
        doc = await db.get(CargoDocument, doc_id)
        if doc and doc.data_json:
            doc_data = json.loads(doc.data_json)

    # Determine current port based on leg status
    if leg.ata:
        current_port = leg.arrival_port.name if leg.arrival_port else ""
    else:
        current_port = leg.departure_port.name if leg.departure_port else ""

    # Pre-fill from leg data
    prefill = {
        "vessel_name": leg.vessel.name if leg.vessel else "",
        "voyage_no": leg.leg_code or "",
        "port": (leg.departure_port.name if leg.departure_port else "") + " / " + (leg.arrival_port.name if leg.arrival_port else ""),
        "port_dep": leg.departure_port.name if leg.departure_port else "",
        "port_dep_locode": leg.departure_port.locode if leg.departure_port else "",
        "port_arr": leg.arrival_port.name if leg.arrival_port else "",
        "port_arr_locode": leg.arrival_port.locode if leg.arrival_port else "",
        "current_port": current_port,
        "etd": leg.etd.strftime("%Y-%m-%d") if leg.etd else "",
        "eta": leg.eta.strftime("%Y-%m-%d") if leg.eta else "",
        "date_today": date.today().strftime("%Y-%m-%d"),
    }

    # Load SOF events for this leg (for pre-filling)
    sof_result = await db.execute(
        select(SofEvent).where(SofEvent.leg_id == leg_id)
        .order_by(SofEvent.event_date.asc().nulls_last(), SofEvent.id.asc())
    )
    sof_events = sof_result.scalars().all()

    # Load embarked crew for this vessel
    crew_result = await db.execute(
        select(CrewAssignment).options(selectinload(CrewAssignment.member))
        .where(
            CrewAssignment.vessel_id == leg.vessel_id,
            CrewAssignment.status == "active",
        ).order_by(CrewAssignment.member_id)
    )
    crew_onboard = crew_result.scalars().all()

    doc_label = next((l for c, l in CARGO_DOC_TYPES if c == doc_type), doc_type)

    # Load attachments for this document
    doc_attachments = []
    if doc:
        att_result = await db.execute(
            select(CargoDocumentAttachment)
            .where(CargoDocumentAttachment.document_id == doc.id)
            .order_by(CargoDocumentAttachment.created_at.asc())
        )
        doc_attachments = att_result.scalars().all()

    # Port timezone
    _ref_port = leg.arrival_port if leg.ata else leg.departure_port
    _port_tz = get_port_timezone(_ref_port.country_code, _ref_port.zone_code) if _ref_port else "UTC"
    _port_tz_label = _ref_port.name if _ref_port else "Port"

    return templates.TemplateResponse("onboard/doc_form.html", {
        "request": request, "user": user,
        "leg": leg, "doc_type": doc_type, "doc_label": doc_label,
        "doc": doc, "doc_data": doc_data, "prefill": prefill,
        "sof_events": sof_events,
        "crew_onboard": crew_onboard,
        "doc_attachments": doc_attachments,
        "tz_choices": TIMEZONE_CHOICES,
        "port_timezone": _port_tz,
        "port_tz_label": _port_tz_label,
        "port_tz_offset": utc_offset_label(_port_tz),
        "active_module": "captain",
        "lang": user.language or "fr",
    })


@router.post("/doc/{leg_id}/{doc_type}/save", response_class=HTMLResponse)
async def cargo_doc_save(
    leg_id: int, doc_type: str, request: Request,
    doc_id: Optional[int] = Form(None),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Save cargo document form data."""
    form = await request.form()
    data = {k: v for k, v in form.items() if k not in ("doc_id",)}

    doc_label = next((l for c, l in CARGO_DOC_TYPES if c == doc_type), doc_type)

    if doc_id:
        doc = await db.get(CargoDocument, doc_id)
        if doc:
            doc.data_json = json.dumps(data, default=str)
            doc.updated_at = func.now()
    else:
        doc = CargoDocument(
            leg_id=leg_id,
            doc_type=doc_type,
            title=doc_label,
            data_json=json.dumps(data, default=str),
            created_by=user.full_name,
        )
        db.add(doc)
        await db.flush()

        # ─── Add SOF event on document creation ───
        sof_evt = SofEvent(
            leg_id=leg_id,
            event_type="CUSTOM",
            event_label=f"📄 Document créé : {doc_label}",
            event_date=date.today(),
            event_time=datetime.now().strftime("%H:%M"),
            remarks=f"Créé par {user.full_name} — /onboard/doc/{doc.id}/export/pdf",
            created_by=user.full_name,
        )
        db.add(sof_evt)

    await db.flush()

    # Redirect back to the document form so export buttons become available
    saved_doc_id = doc.id if doc else None
    if not saved_doc_id:
        # Get the last inserted doc
        last = await db.execute(
            select(CargoDocument).where(
                CargoDocument.leg_id == leg_id, CargoDocument.doc_type == doc_type
            ).order_by(CargoDocument.id.desc())
        )
        saved = last.scalar_one_or_none()
        saved_doc_id = saved.id if saved else None

    return RedirectResponse(url=f"/onboard/doc/{leg_id}/{doc_type}?doc_id={saved_doc_id}", status_code=303)


# ═══════════════════════════════════════════════════════════════
#  ESCALE CLOSURE WORKFLOW
# ═══════════════════════════════════════════════════════════════

@router.post("/closure/{leg_id}/submit-review", response_class=HTMLResponse)
async def closure_submit_review(
    leg_id: int, request: Request,
    closure_notes: str = Form(""),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Submit escale for captain review — transition from 'open' to 'review'."""
    from app.permissions import can_modify
    if not can_modify(user, "captain"):
        raise HTTPException(403)
    leg = await db.get(Leg, leg_id)
    if not leg:
        raise HTTPException(404)
    if leg.closure_status not in (None, "open"):
        raise HTTPException(400, detail="L'escale n'est pas dans un état permettant la soumission.")

    leg.closure_status = "review"
    leg.closure_notes = closure_notes.strip() if closure_notes else None
    leg.closure_reviewed_by = user.full_name
    leg.closure_reviewed_at = datetime.now(timezone.utc)
    await db.flush()

    # SOF event
    sof = SofEvent(
        leg_id=leg_id, event_type="CUSTOM",
        event_label="📋 Escale soumise pour revue — Closure review submitted",
        event_date=date.today(),
        event_time=datetime.now().strftime("%H:%M"),
        remarks=f"Par {user.full_name}",
        created_by=user.full_name,
    )
    db.add(sof)
    await log_activity(db, user, "captain", "closure_review", "Leg", leg_id, f"Soumission revue escale {leg.leg_code}")
    await db.flush()

    url = f"/onboard?vessel={leg.vessel_id}&leg_id={leg_id}"
    if request.headers.get("HX-Request"):
        return HTMLResponse(content="", headers={"HX-Redirect": url})
    return RedirectResponse(url=url, status_code=303)


@router.post("/closure/{leg_id}/approve", response_class=HTMLResponse)
async def closure_approve(
    leg_id: int, request: Request,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Captain approves the escale — transition from 'review' to 'approved'."""
    from app.permissions import can_modify
    if not can_modify(user, "captain"):
        raise HTTPException(403)
    leg = await db.get(Leg, leg_id)
    if not leg:
        raise HTTPException(404)
    if leg.closure_status != "review":
        raise HTTPException(400, detail="L'escale doit être en revue pour être approuvée.")

    leg.closure_status = "approved"
    leg.closure_approved_by = user.full_name
    leg.closure_approved_at = datetime.now(timezone.utc)
    await db.flush()

    sof = SofEvent(
        leg_id=leg_id, event_type="CUSTOM",
        event_label="✅ Escale approuvée par le commandant — Captain approved",
        event_date=date.today(),
        event_time=datetime.now().strftime("%H:%M"),
        remarks=f"Approuvé par {user.full_name}",
        created_by=user.full_name,
    )
    db.add(sof)
    await log_activity(db, user, "captain", "closure_approve", "Leg", leg_id, f"Approbation escale {leg.leg_code}")
    await db.flush()

    url = f"/onboard?vessel={leg.vessel_id}&leg_id={leg_id}"
    if request.headers.get("HX-Request"):
        return HTMLResponse(content="", headers={"HX-Redirect": url})
    return RedirectResponse(url=url, status_code=303)


@router.post("/closure/{leg_id}/reopen", response_class=HTMLResponse)
async def closure_reopen(
    leg_id: int, request: Request,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Reopen an escale that was in review or approved — back to 'open'."""
    from app.permissions import can_modify
    if not can_modify(user, "captain"):
        raise HTTPException(403)
    leg = await db.get(Leg, leg_id)
    if not leg:
        raise HTTPException(404)
    if leg.closure_status not in ("review", "approved"):
        raise HTTPException(400, detail="L'escale ne peut pas être rouverte depuis cet état.")

    old_status = leg.closure_status
    leg.closure_status = "open"
    leg.closure_reviewed_by = None
    leg.closure_reviewed_at = None
    leg.closure_approved_by = None
    leg.closure_approved_at = None
    await db.flush()

    sof = SofEvent(
        leg_id=leg_id, event_type="CUSTOM",
        event_label=f"🔓 Escale rouverte (depuis {old_status}) — Closure reopened",
        event_date=date.today(),
        event_time=datetime.now().strftime("%H:%M"),
        remarks=f"Par {user.full_name}",
        created_by=user.full_name,
    )
    db.add(sof)
    await log_activity(db, user, "captain", "closure_reopen", "Leg", leg_id, f"Réouverture escale {leg.leg_code}")
    await db.flush()

    url = f"/onboard?vessel={leg.vessel_id}&leg_id={leg_id}"
    if request.headers.get("HX-Request"):
        return HTMLResponse(content="", headers={"HX-Redirect": url})
    return RedirectResponse(url=url, status_code=303)


@router.post("/closure/{leg_id}/lock", response_class=HTMLResponse)
async def closure_lock(
    leg_id: int, request: Request,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Final lock after approval — generates PDF and locks escale."""
    from app.permissions import can_modify
    if not can_modify(user, "captain"):
        raise HTTPException(403)
    leg_result = await db.execute(
        select(Leg).options(
            selectinload(Leg.vessel),
            selectinload(Leg.departure_port),
            selectinload(Leg.arrival_port),
        ).where(Leg.id == leg_id)
    )
    leg = leg_result.scalar_one_or_none()
    if not leg:
        raise HTTPException(404)
    if leg.closure_status != "approved":
        raise HTTPException(400, detail="L'escale doit être approuvée avant le verrouillage.")

    # Generate closure PDF
    pdf_path = await _generate_closure_pdf(db, leg, user)

    leg.closure_status = "locked"
    leg.closure_pdf_path = pdf_path
    leg.status = "completed"
    await db.flush()

    sof = SofEvent(
        leg_id=leg_id, event_type="CUSTOM",
        event_label="🔒 Escale clôturée et verrouillée — Port call closed & locked",
        event_date=date.today(),
        event_time=datetime.now().strftime("%H:%M"),
        remarks=f"PDF généré — Verrouillé par {user.full_name}",
        created_by=user.full_name,
    )
    db.add(sof)
    await log_activity(db, user, "captain", "closure_lock", "Leg", leg_id, f"Clôture définitive escale {leg.leg_code}")
    await db.flush()

    url = f"/onboard?vessel={leg.vessel_id}&leg_id={leg_id}"
    if request.headers.get("HX-Request"):
        return HTMLResponse(content="", headers={"HX-Redirect": url})
    return RedirectResponse(url=url, status_code=303)


@router.get("/closure/{leg_id}/pdf")
async def closure_download_pdf(
    leg_id: int,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Download the closure PDF."""
    from app.permissions import can_view
    if not can_view(user, "captain"):
        raise HTTPException(403)
    leg = await db.get(Leg, leg_id)
    if not leg or not leg.closure_pdf_path:
        raise HTTPException(404, detail="PDF de clôture non disponible.")
    import os
    if not os.path.exists(leg.closure_pdf_path):
        raise HTTPException(404, detail="Fichier PDF introuvable.")
    from fastapi.responses import FileResponse
    return FileResponse(
        leg.closure_pdf_path,
        filename=f"closure_{leg.leg_code}.pdf",
        media_type="application/pdf",
    )


async def _generate_closure_pdf(db: AsyncSession, leg, user) -> str:
    """Generate a closure summary PDF for the escale."""
    import os
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    # Fetch data
    sof_result = await db.execute(
        select(SofEvent).where(SofEvent.leg_id == leg.id)
        .order_by(SofEvent.event_date.asc().nulls_last(), SofEvent.event_time.asc().nulls_last())
    )
    sof_events = sof_result.scalars().all()

    doc_result = await db.execute(
        select(CargoDocument).where(CargoDocument.leg_id == leg.id).order_by(CargoDocument.created_at)
    )
    cargo_docs = doc_result.scalars().all()

    att_result = await db.execute(
        select(OnboardAttachment).where(OnboardAttachment.leg_id == leg.id).order_by(OnboardAttachment.created_at)
    )
    attachments_list = att_result.scalars().all()

    orders_result = await db.execute(select(Order).where(Order.leg_id == leg.id))
    orders = orders_result.scalars().all()

    crew_result = await db.execute(
        select(CrewAssignment).options(selectinload(CrewAssignment.member))
        .where(CrewAssignment.vessel_id == leg.vessel_id, CrewAssignment.status == "active")
    )
    crew = crew_result.scalars().all()

    # PDF generation
    os.makedirs("app/uploads/closure", exist_ok=True)
    pdf_path = f"app/uploads/closure/closure_{leg.leg_code}_{int(datetime.now().timestamp())}.pdf"

    doc = SimpleDocTemplate(pdf_path, pagesize=A4, leftMargin=20*mm, rightMargin=20*mm, topMargin=15*mm, bottomMargin=15*mm)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="TitleBlue", fontName="Helvetica-Bold", fontSize=16, textColor=colors.HexColor("#095561"), alignment=TA_CENTER, spaceAfter=12))
    styles.add(ParagraphStyle(name="SectionH", fontName="Helvetica-Bold", fontSize=12, textColor=colors.HexColor("#095561"), spaceBefore=14, spaceAfter=6))
    styles.add(ParagraphStyle(name="Small", fontName="Helvetica", fontSize=9, textColor=colors.HexColor("#555555")))

    elements = []

    # Header
    elements.append(Paragraph(f"RAPPORT DE CLÔTURE D'ESCALE", styles["TitleBlue"]))
    elements.append(Paragraph(f"TOWT — Transoceanic Wind Transport", styles["Small"]))
    elements.append(Spacer(1, 8))

    # Leg info table
    vessel_name = leg.vessel.name if leg.vessel else "—"
    dep_port = leg.departure_port.name if leg.departure_port else "—"
    arr_port = leg.arrival_port.name if leg.arrival_port else "—"
    info_data = [
        ["Leg Code", leg.leg_code, "Navire", vessel_name],
        ["Port départ", dep_port, "Port arrivée", arr_port],
        ["ETD", leg.etd.strftime("%d/%m/%Y %H:%M") if leg.etd else "—", "ETA", leg.eta.strftime("%d/%m/%Y %H:%M") if leg.eta else "—"],
        ["ATD", leg.atd.strftime("%d/%m/%Y %H:%M") if leg.atd else "—", "ATA", leg.ata.strftime("%d/%m/%Y %H:%M") if leg.ata else "—"],
    ]
    info_table = Table(info_data, colWidths=[80, 140, 80, 140])
    info_table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTNAME", (2, 0), (2, -1), "Helvetica-Bold"),
        ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#095561")),
        ("TEXTCOLOR", (2, 0), (2, -1), colors.HexColor("#095561")),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e0e0e0")),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f0f7f8")),
        ("BACKGROUND", (2, 0), (2, -1), colors.HexColor("#f0f7f8")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    elements.append(info_table)
    elements.append(Spacer(1, 10))

    # Cargo summary
    elements.append(Paragraph("CARGO", styles["SectionH"]))
    total_palettes = sum(o.quantity_palettes or 0 for o in orders)
    total_weight = sum(o.total_weight or 0 for o in orders)
    cargo_data = [["Commandes", "Palettes", "Poids total"]]
    cargo_data.append([str(len(orders)), str(total_palettes), f"{total_weight:.1f} t"])
    cargo_table = Table(cargo_data, colWidths=[140, 140, 160])
    cargo_table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e0e0e0")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f0f7f8")),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    elements.append(cargo_table)

    # SOF summary
    if sof_events:
        elements.append(Paragraph("STATEMENT OF FACTS", styles["SectionH"]))
        sof_data = [["Événement", "Date", "Heure", "Remarques"]]
        for evt in sof_events:
            sof_data.append([
                evt.event_label[:60],
                evt.event_date.strftime("%d/%m/%Y") if evt.event_date else "",
                evt.event_time or "",
                (evt.remarks or "")[:40],
            ])
        sof_table = Table(sof_data, colWidths=[180, 70, 50, 140])
        sof_table.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e0e0e0")),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f0f7f8")),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        elements.append(sof_table)

    # Cargo documents
    if cargo_docs:
        elements.append(Paragraph("DOCUMENTS CARGO", styles["SectionH"]))
        for cd in cargo_docs:
            doc_type_label = dict(CARGO_DOC_TYPES).get(cd.doc_type, cd.doc_type)
            elements.append(Paragraph(f"• {doc_type_label} — créé par {cd.created_by or '—'} le {cd.created_at.strftime('%d/%m/%Y %H:%M') if cd.created_at else '—'}", styles["Small"]))

    # Attachments
    if attachments_list:
        elements.append(Paragraph("PIÈCES JOINTES", styles["SectionH"]))
        for att in attachments_list:
            elements.append(Paragraph(f"• [{att.category}] {att.title} — {att.filename} ({att.uploaded_by})", styles["Small"]))

    # Crew
    if crew:
        elements.append(Paragraph("ÉQUIPAGE", styles["SectionH"]))
        crew_data = [["Nom", "Rôle"]]
        for ca in crew:
            crew_data.append([f"{ca.member.first_name} {ca.member.last_name}", ca.role.replace("_", " ").capitalize()])
        crew_table = Table(crew_data, colWidths=[220, 220])
        crew_table.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e0e0e0")),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f0f7f8")),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]))
        elements.append(crew_table)

    # Closure info
    elements.append(Spacer(1, 16))
    elements.append(Paragraph("VALIDATION", styles["SectionH"]))
    now_str = datetime.now().strftime("%d/%m/%Y %H:%M")
    closure_info = [
        f"Revue par : {leg.closure_reviewed_by or '—'} — {leg.closure_reviewed_at.strftime('%d/%m/%Y %H:%M') if leg.closure_reviewed_at else '—'}",
        f"Approuvé par : {leg.closure_approved_by or '—'} — {leg.closure_approved_at.strftime('%d/%m/%Y %H:%M') if leg.closure_approved_at else '—'}",
        f"Verrouillé par : {user.full_name} — {now_str}",
    ]
    for line in closure_info:
        elements.append(Paragraph(line, styles["Small"]))
    if leg.closure_notes:
        elements.append(Spacer(1, 6))
        elements.append(Paragraph(f"Notes : {leg.closure_notes}", styles["Small"]))

    doc.build(elements)
    return pdf_path
